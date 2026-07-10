# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x39, 0x64)
GREEN = RGBColor(0x00, 0x6A, 0x4E)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexc); tcPr.append(shd)


def ct(cell, text, bold=False, white=False, size=9, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color: r.font.color.rgb = color


def h1(t):
    p = doc.add_heading(level=1)
    r = p.add_run(t); r.font.color.rgb = NAVY; r.font.size = Pt(15); return p


def h2(t):
    p = doc.add_heading(level=2)
    r = p.add_run(t); r.font.color.rgb = GREEN; r.font.size = Pt(12.5); return p


def h3(t):
    p = doc.add_heading(level=3)
    r = p.add_run(t); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33); r.font.size = Pt(11); return p


def para(t, bold=False, italic=False, size=10.5, color=None, after=6):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p


def bullet(t, pre=None):
    p = doc.add_paragraph(style='List Bullet')
    if pre:
        rr = p.add_run(pre); rr.bold = True
    p.add_run(t); return p


def feat_table(rows):
    t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
    for i, htxt in enumerate(['Fitur', 'Deskripsi']):
        shade(t.rows[0].cells[i], '006A4E'); ct(t.rows[0].cells[i], htxt, bold=True, white=True)
    for name, desc in rows:
        c = t.add_row().cells
        ct(c[0], name, bold=True, size=9); ct(c[1], desc, size=9)
        c[0].width = Inches(2.1); c[1].width = Inches(4.6)
    doc.add_paragraph()


def steps(items):
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number'); p.add_run(s)


def test_table(rows):
    t = doc.add_table(rows=1, cols=4); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    heads = ['ID', 'Skenario', 'Langkah', 'Hasil yang Diharapkan']
    widths = [Inches(0.5), Inches(1.5), Inches(2.6), Inches(2.1)]
    for i, htxt in enumerate(heads):
        shade(t.rows[0].cells[i], '1F3964'); ct(t.rows[0].cells[i], htxt, bold=True, white=True, size=9)
    for tid, sc, lg, exp in rows:
        c = t.add_row().cells
        ct(c[0], tid, bold=True, size=8.5); ct(c[1], sc, size=8.5)
        ct(c[2], lg, size=8.5); ct(c[3], exp, size=8.5)
        for i, w in enumerate(widths): c[i].width = w
    doc.add_paragraph()


# ============================ COVER ============================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('DAFTAR FITUR, USER GUIDE & SKENARIO TEST'); r.bold = True; r.font.size = Pt(19); r.font.color.rgb = NAVY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Kustomisasi Property Management Odoo 19\n(Companion Modules untuk addon rental_management)')
r.font.size = Pt(13); r.bold = True; r.font.color.rgb = GREEN
doc.add_paragraph()
meta = doc.add_table(rows=5, cols=2); meta.style = 'Table Grid'
for i, (k, v) in enumerate([
    ('Dokumen', 'Feature List, User Guide & Test Scenarios'),
    ('Lingkup', '26 modul companion + integrasi modul standar Odoo 19'),
    ('Basis', 'addon rental_management v3.3.9 (TechKhedut) – Odoo 19'),
    ('Disusun oleh', 'System Analyst / Functional & Odoo Developer'),
    ('Tanggal', '19 Juni 2026'),
]):
    shade(meta.rows[i].cells[0], 'E7ECF5'); ct(meta.rows[i].cells[0], k, bold=True, size=10)
    ct(meta.rows[i].cells[1], v, size=10)
    meta.rows[i].cells[0].width = Inches(1.7); meta.rows[i].cells[1].width = Inches(5.0)
doc.add_page_break()

# ============================ 1. OVERVIEW ============================
h1('1. Ikhtisar & Daftar Modul')
para('Dokumen ini memuat detail fitur, panduan penggunaan, dan skenario pengujian untuk '
     'seluruh modul kustom yang dikembangkan di atas addon pihak ketiga rental_management '
     '(TechKhedut). Setiap modul bersifat companion (tidak menambal modul berlisensi) dan '
     'terintegrasi dengan modul standar Odoo 19.')

ov = doc.add_table(rows=1, cols=3); ov.style = 'Table Grid'
for i, htxt in enumerate(['Modul', 'Fungsi Utama', 'Depends (modul standar)']):
    shade(ov.rows[0].cells[i], '1F3964'); ct(ov.rows[0].cells[i], htxt, bold=True, white=True)
modules_ov = [
    ('rental_management_financial_report', 'Owners Statement (9 laporan), Trust Accounting, Owner Remittance, Budget, Security Deposit, integrasi Analytic', 'account, (analytic)'),
    ('rental_management_gto_meter', 'GTO / Revenue Sharing & Meter Management (recharge utilitas)', 'account, product'),
    ('rental_management_guarantee', 'Bank/Insurance Guarantee + alert kedaluwarsa', 'mail'),
    ('rental_management_casual_leasing', 'Sewa jangka pendek (booth/kios/pop-up)', 'account, product'),
    ('rental_management_handover', 'Move-in / Fit-out / Move-out (checklist & dokumen)', 'mail'),
    ('rental_management_portal', 'Portal mandiri tenant (/my)', 'portal'),
    ('rental_management_coretax', 'e-Faktur & SPT CORETAX (11 jenis ekspor XML)', 'account'),
    ('rental_management_purchase', 'Procurement properti (PO ↔ properti/kontrak/maintenance)', 'purchase, account'),
    ('rental_management_crm', 'Pipeline leasing (lead → kontrak)', 'crm'),
    ('rental_management_project', 'Fit-out sebagai Project & Task', 'project, handover'),
    ('rental_management_documents', 'Folder Documents per properti (Enterprise)', 'documents'),
    ('rental_management_asset', 'Fixed Asset: penyusutan (linear/declining) + revaluation + sync CORETAX L9', 'account'),
    ('rental_management_document_ce', 'Register lampiran per properti (fallback Community untuk Documents)', 'rental_management'),
    ('rental_management_owner_portal', 'Portal mandiri pemilik (/my/properties + remittance)', 'portal'),
    ('rental_management_cam', 'CAM / Service Charge: pool biaya, apportion per area, tagih', 'account'),
    ('rental_management_rent_escalation', 'Eskalasi sewa berkala (fixed %/amount) + log', 'rental_management'),
    ('rental_management_dashboard', 'Dashboard KPI manajemen (NOI, arrears, collection, expiry)', 'account'),
    ('rental_management_dunning', 'Dunning ladder otomatis (reminder email + late fee)', 'account, mail'),
    ('rental_management_insurance', 'Register polis asuransi properti + alert kedaluwarsa', 'mail'),
    ('rental_management_lease_expiry', 'Daftar lease expiry (30/90 hari) + reminder renewal', 'rental_management, mail'),
    ('rental_management_vacancy', 'Board vacancy/availability per status', 'rental_management'),
    ('rental_management_ppm', 'Preventive maintenance plan auto-generate maintenance.request', 'rental_management'),
    ('rental_management_parking', 'Register car-park bay, alokasi tenant, invoice parkir', 'account'),
    ('rental_management_valuation', 'Register valuasi pasar berkala per properti', 'rental_management'),
    ('rental_management_access', 'Register kartu akses/parkir + log kunjungan tamu', 'rental_management'),
    ('rental_management_esign', 'Tracking permintaan tanda tangan kontrak (non-Enterprise)', 'rental_management, mail'),
]
for n, f, d in modules_ov:
    c = ov.add_row().cells
    ct(c[0], n, bold=True, size=8.5); ct(c[1], f, size=8.5); ct(c[2], d, size=8.5)
    c[0].width = Inches(2.3); c[1].width = Inches(3.0); c[2].width = Inches(1.4)
doc.add_paragraph()

h2('1.1 Prasyarat & Urutan Instalasi')
bullet('addon rental_management (TechKhedut) harus sudah terpasang lebih dulu.', pre='Wajib: ')
bullet('Aktifkan modul standar: Accounting, Purchase, Project, CRM, Portal (Documents = Enterprise).', pre='Aktifkan: ')
bullet('Install rental_management_financial_report dahulu (banyak modul lain memanfaatkan field property_manual_id & analytic untuk mengalir ke laporan).', pre='Disarankan: ')
bullet('Modul lain dapat diinstal independen sesuai kebutuhan; semua memakai pengecekan keberadaan field agar decoupled.', pre='Catatan: ')

h2('1.2 Konvensi Skenario Test')
para('Setiap modul memiliki tabel skenario test dengan kolom: ID, Skenario, Langkah, dan Hasil '
     'yang Diharapkan. Prasyarat umum: user login dengan hak akses memadai (Accounting/User atau '
     'Manager), minimal 1 properti dan 1 kontrak (tenancy.details) aktif tersedia.', italic=True, color=GREY)
doc.add_page_break()

# ============================ MODULE SECTIONS ============================

# ---- FINANCIAL REPORT ----
h1('2. Modul: Financial Report (Owners Statement)')
para('Nama teknis: rental_management_financial_report — Menu: Properties → Financial Reports.', italic=True, color=GREY)
h2('2.1 Daftar Fitur')
feat_table([
    ('Owners Statement (PDF)', 'Satu dokumen berisi 9 laporan: Performance Summary (Accrual + Cash & Trust), Income & Expenditure (Accrual), Receipts & Payments (Cash), Tenant Balances, Aged Arrears, Payment Details, Trial Balance, Balance Sheet, GST Reconciliation.'),
    ('Multi-owner & Ownership %', 'Banyak pemilik per properti dengan persentase, Property Manager, telepon, faks.'),
    ('Report Categories', 'Pemetaan akun GL ke section/sub-group laporan (struktur MR-style).'),
    ('Property Budget', 'Anggaran per akun per bulan → Actual vs Budget vs Variance vs %Var.'),
    ('Trust Accounting', 'Trust Bank Account; Opening/Closing Trust Balance, Available for Remittance.'),
    ('Owner Remittance', 'Pembagian dana trust ke pemilik per %; posting Dr Owners Remittance / Cr Trust Bank.'),
    ('Security Deposit', 'Deposit sebagai liabilitas (bukan income), deduction/refund, saldo → kolom Sec Dep Bal.'),
    ('Integrasi Analytic', 'Saat posting, item income/expense ber-tag properti otomatis dicap analytic account properti.'),
    ('Budget Approval Workflow', 'Property Budget: Draft → To Approve → Approved (approver + tanggal); Approve dibatasi grup Accounting Manager.'),
    ('Multi-currency Remittance', 'Owner Remittance dapat memilih mata uang; posting dikonversi ke mata uang perusahaan pada kurs tanggal remittance (amount_currency).'),
])
h2('2.2 User Guide')
steps([
    'Buka properti → tab "Owners & Financial": isi Owners + %, Property Manager, dan klik "Create Analytic Account".',
    'Isi konfigurasi Trust Accounting (Trust Bank Account, Owners Remittance Account, Remittance Journal) dan Security Deposit (akun liabilitas, akun forfeiture, journal).',
    'Buka Accounting → Chart of Accounts, set "Property Report Category" pada akun income/expense terkait (atau via menu Report Categories).',
    'Menu Financial Reports → Property Budgets: buat anggaran per akun/bulan (opsional "Generate Monthly Lines").',
    'Menu Financial Reports → Security Deposits / Owner Remittances: catat deposit & remittance, lalu Post.',
    'Menu Financial Reports → Owners Statement: pilih properti, periode, awal tahun fiskal → "Print Owners Statement".',
])
h2('2.3 Skenario Test')
test_table([
    ('FR-01', 'Cetak Owners Statement dasar', 'Pastikan ada invoice/bill posted ber-link properti. Jalankan wizard Owners Statement untuk periode berjalan.', 'PDF tampil dengan 9 section; angka Income/Expense terisi dari GL.'),
    ('FR-02', 'Actual vs Budget', 'Buat Property Budget per akun untuk bulan uji. Cetak ulang Owners Statement.', 'Kolom Budget & Variance & %Var pada I&E dan Performance Summary terisi sesuai budget.'),
    ('FR-03', 'Multi-owner & remittance', 'Isi 2 owner (60/40). Buat Owner Remittance → Compute from Owners → Post.', 'Alokasi terbagi 60/40; jurnal Dr Remittance/Cr Trust terbuat; Less Remittances muncul di Performance Summary (Cash).'),
    ('FR-04', 'Trust roll-forward', 'Set Trust Bank Account; pastikan ada penerimaan & pembayaran ter-rekonsiliasi.', 'Opening + Net Cash − Remittances = Closing Trust Balance konsisten.'),
    ('FR-05', 'Security Deposit liability', 'Buat Security Deposit, klik Mark as Held; tambah baris refund; Post Deductions/Refunds.', 'Balance berkurang; kolom Sec Dep Bal di Tenant Balances menampilkan saldo.'),
    ('FR-06', 'Integrasi Analytic', 'Set analytic account properti. Posting sebuah invoice ber-link properti. Buka tombol "Analytic Items".', 'Baris income/expense memiliki analytic_distribution ke analytic properti; muncul di Analytic Items.'),
    ('FR-07', 'Aged Arrears', 'Biarkan beberapa invoice jatuh tempo (umur berbeda). Cetak laporan.', 'Tunggakan terbagi ke bucket Current/1/2/3/4+ bulan dengan total benar.'),
])

# ---- GTO & METER ----
h1('3. Modul: GTO & Meter Management')
para('Nama teknis: rental_management_gto_meter — Menu: Properties → Leasing Operations.', italic=True, color=GREY)
h2('3.1 Daftar Fitur')
feat_table([
    ('GTO / Revenue Sharing', 'Konfigurasi di kontrak (Higher-of / Base+Overage / Pure %), Turnover %, Breakpoint, produk GTO.'),
    ('Turnover Declaration', 'Deklarasi omzet tenant per periode; hitung otomatis percentage/overage rent; buat invoice.'),
    ('Meter Management', 'Meter listrik/air/gas per properti dengan produk recharge & tarif.'),
    ('Meter Reading', 'Pembacaan berurutan (previous auto), konsumsi & jumlah otomatis, buat invoice recharge ke tenant.'),
])
h2('3.2 User Guide')
steps([
    'Buka kontrak (tenancy) → tab "GTO / Revenue Sharing": centang GTO, pilih tipe, isi Turnover % (dan Breakpoint bila Base+Overage).',
    'Leasing Operations → GTO Turnovers: buat deklarasi (periode + omzet), cek Billable, klik Create Invoice.',
    'Leasing Operations → Meters: daftarkan meter (properti, produk recharge, tarif, satuan).',
    'Leasing Operations → Meter Readings: input pembacaan baru, cek konsumsi & amount, klik Create Recharge Invoice.',
])
h2('3.3 Skenario Test')
test_table([
    ('GT-01', 'GTO Higher-of', 'Kontrak GTO higher_of, %=5, base rent=10jt. Turnover=300jt. Compute.', 'percentage_rent=15jt; billable=5jt (15jt−10jt); invoice terbuat sejumlah 5jt.'),
    ('GT-02', 'GTO Base+Overage', 'Tipe base_plus, %=5, breakpoint=200jt, turnover=300jt.', 'billable = (300−200)jt×5% = 5jt.'),
    ('GT-03', 'GTO Pure %', 'Tipe pure, %=8, turnover=100jt.', 'billable = 8jt.'),
    ('MT-01', 'Meter reading & recharge', 'Meter tarif=1.500/kWh, previous=100. Input current=350. Create Recharge Invoice.', 'consumption=250; amount=375.000; invoice out_invoice ke tenant qty 250 × 1.500.'),
    ('MT-02', 'Previous reading otomatis', 'Buat reading kedua pada meter yang sama.', 'previous_reading terisi otomatis dari current reading terakhir.'),
    ('GT-04', 'Aliran ke laporan', 'Posting invoice GTO/meter, cetak Owners Statement.', 'Pendapatan muncul di Income (Tenant Recharge / Rental) Owners Statement.'),
])

# ---- GUARANTEE ----
h1('4. Modul: Bank & Insurance Guarantee')
para('Nama teknis: rental_management_guarantee — Menu: Properties → Tenant Guarantees.', italic=True, color=GREY)
h2('4.1 Daftar Fitur')
feat_table([
    ('Registrasi Jaminan', 'Bank Guarantee / Insurance Bond / Security Guarantee, nomor, penerbit, nilai, tanggal terbit & expiry.'),
    ('Lifecycle', 'Draft → Active → Expired / Released / Claimed.'),
    ('Indikator Expiry', 'Days-to-expiry + pewarnaan list (merah=expired, oranye=expiring soon).'),
    ('Alert Otomatis', 'Cron harian: auto-expire jaminan lewat tanggal & jadwalkan reminder activity sebelum expiry.'),
])
h2('4.2 User Guide')
steps([
    'Properties → Tenant Guarantees → New: pilih kontrak, tipe, nomor, penerbit, nilai, tanggal terbit & expiry, lead reminder (default 30 hari).',
    'Klik Activate. Pantau kolom Days to Expiry / filter "Expiring Soon".',
    'Saat selesai: klik Release atau Claim sesuai kondisi.',
])
h2('4.3 Skenario Test')
test_table([
    ('GU-01', 'Buat & aktifkan jaminan', 'Buat guarantee expiry 60 hari ke depan, Activate.', 'State=Active; days_to_expiry≈60.'),
    ('GU-02', 'Expiring soon', 'Set expiry dalam 10 hari (reminder 30). Jalankan cron / tunggu cron.', 'is_expiring=true; activity reminder terjadwal pada responsible.'),
    ('GU-03', 'Auto-expire', 'Set expiry kemarin. Jalankan cron.', 'State berubah ke Expired; ada log pesan.'),
    ('GU-04', 'Release/Claim', 'Pada jaminan Active, klik Release lalu pada lain Claim.', 'State berubah sesuai aksi.'),
])

# ---- CASUAL LEASING ----
h1('5. Modul: Casual Leasing')
para('Nama teknis: rental_management_casual_leasing — Menu: Properties → Casual Leases.', italic=True, color=GREY)
h2('5.1 Daftar Fitur')
feat_table([
    ('Casual Lease', 'Properti + space/lokasi, customer, periode, tarif Per Hari/Per Minggu/Fixed.'),
    ('Perhitungan Otomatis', 'Quantity & total terhitung (hari, ceil(hari/7) minggu, atau 1).'),
    ('Invoice 1-klik', 'Buat invoice pelanggan; ber-tag properti bila modul financial terpasang.'),
    ('Lifecycle', 'Draft → Confirmed → Active → Done / Cancelled.'),
])
h2('5.2 User Guide')
steps([
    'Properties → Casual Leases → New: pilih properti, space, customer, periode, tipe tarif & tarif.',
    'Periksa Quantity & Total otomatis. Klik Confirm.',
    'Klik Create Invoice untuk menagih customer.',
])
h2('5.3 Skenario Test')
test_table([
    ('CL-01', 'Tarif per hari', 'Periode 5 hari, Per Day, tarif=200rb.', 'duration_days=5; total=1jt.'),
    ('CL-02', 'Tarif per minggu', 'Periode 10 hari, Per Week, tarif=1jt.', 'quantity=ceil(10/7)=2; total=2jt.'),
    ('CL-03', 'Invoice', 'Confirm lalu Create Invoice.', 'out_invoice terbuat sejumlah total; state=Active; ber-link properti.'),
])

# ---- HANDOVER ----
h1('6. Modul: Move-in / Fit-out / Move-out (Handover)')
para('Nama teknis: rental_management_handover — Menu: Properties → Handovers.', italic=True, color=GREY)
h2('6.1 Daftar Fitur')
feat_table([
    ('Handover Record', 'Tipe Move-in / Fit-out / Move-out terhubung kontrak.'),
    ('Field per Tipe', 'Fit-out: kontraktor, bond, tanggal; Move-out: make-good, kunci dikembalikan; condition report.'),
    ('Checklist + Progress', 'Template default 1-klik per tipe + progress bar % selesai.'),
    ('Dokumen', 'Lampiran via chatter (condition report, gambar, sertifikat, izin).'),
])
h2('6.2 User Guide')
steps([
    'Properties → Handovers → New: pilih tipe & kontrak, tanggal terjadwal, penanggung jawab.',
    'Klik "Load Default Checklist" untuk mengisi checklist sesuai tipe.',
    'Klik Start; centang item checklist (progress bar naik); lampirkan dokumen via chatter.',
    'Klik Complete saat selesai (actual date terisi).',
])
h2('6.3 Skenario Test')
test_table([
    ('HO-01', 'Move-in checklist', 'Buat handover Move-in, Load Default Checklist.', '5 item checklist Move-in muncul.'),
    ('HO-02', 'Progress', 'Centang 3 dari 6 item.', 'progress=50%.'),
    ('HO-03', 'Lifecycle', 'Start lalu Complete.', 'State=Completed; actual_date terisi otomatis.'),
    ('HO-04', 'Fit-out fields', 'Tipe Fit-out: isi kontraktor & bond.', 'Field fit-out tampil hanya untuk tipe Fit-out.'),
])

# ---- PORTAL ----
h1('7. Modul: Tenant Portal')
para('Nama teknis: rental_management_portal — URL: /my (portal tenant).', italic=True, color=GREY)
h2('7.1 Daftar Fitur')
feat_table([
    ('Kartu Contracts', 'Di portal home /my dengan jumlah kontrak.'),
    ('/my/contracts', 'Daftar kontrak milik tenant (paginasi).'),
    ('/my/contracts/<id>', 'Detail kontrak + tombol ke /my/invoices.'),
    ('Keamanan', 'Record rule: tenant hanya melihat kontraknya; nama properti via stored related.'),
])
h2('7.2 User Guide')
steps([
    'Contacts → pilih kontak tenant → Action → Grant portal access (kirim undangan).',
    'Tenant login ke portal → klik kartu Contracts → lihat daftar & detail kontrak.',
    'Dari detail kontrak, klik "View My Invoices" untuk melihat tagihan.',
])
h2('7.3 Skenario Test')
test_table([
    ('PT-01', 'Akses portal tenant', 'Login sebagai tenant portal yang punya kontrak.', 'Kartu Contracts tampil dengan jumlah benar.'),
    ('PT-02', 'Isolasi data', 'Tenant A mencoba akses kontrak tenant B (ubah id URL).', 'Diarahkan ke /my (akses ditolak).'),
    ('PT-03', 'Detail & invoice', 'Buka detail kontrak, klik View My Invoices.', 'Detail tampil; daftar invoice tenant muncul.'),
])

# ---- CORETAX ----
h1('8. Modul: CORETAX e-Faktur & SPT')
para('Nama teknis: rental_management_coretax — Menu: Properties → CORETAX.', italic=True, color=GREY)
h2('8.1 Daftar Fitur (11 jenis ekspor XML)')
feat_table([
    ('Faktur Keluaran (PK)', 'TaxInvoiceBulk v1.4 dari invoice pelanggan.'),
    ('Retur Faktur Masukan', 'InputTaxInvoiceReturn dari vendor credit note.'),
    ('Lampiran C', 'VAT/STLG dipungut pemungut lain.'),
    ('Pencatatan', 'SimpleBookKeepingBulk dari invoice.'),
    ('L9 Penyusutan/Amortisasi', 'DepreciationAmortization dari register aset (desimal dipertahankan).'),
    ('L3B PPh Pihak Lain', 'OtherParties dari register.'),
    ('L11A Piutang Tak Tertagih / Kredit Kurang Lancar', 'UncollectibleDebtBulk & NonPerforming dari register.'),
    ('L11A Promosi / Entertainment', 'PromotionExpense & EntertainmentExpense dari register.'),
    ('L10A Hubungan Istimewa', 'DeclarationOfTransactionRelatedPartiesBulk dari register.'),
])
h2('8.2 User Guide')
steps([
    'Set CORETAX TIN/NPWP & ID TKU pada kontak Company dan tiap kontak pelanggan (tab CORETAX).',
    'Set CORETAX Type/Unit Code pada produk (sewa = Service / B).',
    'Pada invoice pelanggan, periksa tab CORETAX e-Faktur (TrxCode default 04, dst.).',
    'Untuk lampiran SPT: isi register di Properties → CORETAX → SPT Registers (per Tax Year).',
    'Menu CORETAX → e-Faktur / SPT Export: pilih Export Type, periode/Tax Year → Export XML → unduh file untuk diunggah ke CORETAX.',
])
h2('8.3 Skenario Test')
test_table([
    ('CT-01', 'Ekspor Faktur PK', 'Set TIN company & buyer. Posting invoice. Export Type=Faktur Keluaran, periode.', 'File XML TaxInvoiceBulk terunduh; struktur tag sesuai v1.4; invoice ditandai exported.'),
    ('CT-02', 'Retur PM', 'Buat vendor credit note (in_refund) + isi Original Faktur No. Export Type=Retur PM.', 'InputTaxInvoiceReturn berisi Rows + FooterRow total.'),
    ('CT-03', 'L9 Penyusutan', 'Isi register L9 (depreciation & amortization), Tax Year. Export.', 'DepreciationAmortization dengan ListOfDepreciation/Amortization; nilai desimal utuh.'),
    ('CT-04', 'Validasi tanpa TIN', 'Hapus TIN company, coba Export PK.', 'Muncul peringatan agar set TIN/NPWP.'),
])

# ---- PURCHASE ----
h1('9. Modul: Property Procurement (Purchase)')
para('Nama teknis: rental_management_purchase — Menu: Properties → Property Purchase Orders.', italic=True, color=GREY)
h2('9.1 Daftar Fitur')
feat_table([
    ('Link Properti pada PO', 'Property / Lease Contract / Maintenance Request pada purchase.order.'),
    ('Propagasi ke Vendor Bill', 'Saat buat tagihan dari PO, link properti diwariskan → masuk Owners Statement Payment Details + analytic.'),
    ('PO dari Maintenance', 'Tombol Create Purchase Order dari maintenance request (memakai baris produk & vendor).'),
])
h2('9.2 User Guide')
steps([
    'Purchase → buat RFQ/PO, isi tab "Property" (Property/Contract/Maintenance).',
    'Konfirmasi PO, terima barang, lalu Create Bill — bill otomatis ber-link properti.',
    'Atau dari Maintenance Request: klik "Create Purchase Order".',
])
h2('9.3 Skenario Test')
test_table([
    ('PU-01', 'PO ber-properti → bill', 'Buat PO dengan Property terisi, konfirmasi, Create Bill, post.', 'Vendor bill memiliki property link; muncul di Payment Details & analytic properti.'),
    ('PU-02', 'PO dari maintenance', 'Pada maintenance request dengan vendor & produk, klik Create Purchase Order.', 'PO terbuat dengan baris produk request & maintenance_request_id terisi.'),
])

# ---- CRM ----
h1('10. Modul: Leasing CRM Pipeline')
para('Nama teknis: rental_management_crm — Menu: CRM.', italic=True, color=GREY)
h2('10.1 Daftar Fitur')
feat_table([
    ('Field Leasing', 'Expected move-in, expected rent, lease months pada lead.'),
    ('Create Lease Contract', 'Buka form tenancy baru pre-filled dari lead.'),
    ('Smart Button', 'Lease Contracts untuk kontak lead.'),
    ('Team & Tag', '"Property Leasing" team + tag "Leasing".'),
])
h2('10.2 User Guide')
steps([
    'CRM → buat lead, pilih Property (dari addon) & Customer, isi field leasing.',
    'Geser melalui pipeline. Saat deal, klik "Create Lease Contract".',
    'Lengkapi field wajib pada form tenancy yang terbuka, simpan.',
])
h2('10.3 Skenario Test')
test_table([
    ('CR-01', 'Create Lease Contract', 'Pada lead dengan partner & property, klik Create Lease Contract.', 'Form tenancy.details terbuka dengan property & tenant ter-default.'),
    ('CR-02', 'Smart button kontrak', 'Lead dengan partner yang punya kontrak.', 'Smart button menampilkan jumlah & daftar kontrak.'),
])

# ---- PROJECT ----
h1('11. Modul: Fit-out Projects')
para('Nama teknis: rental_management_project — pada form Handover (Fit-out).', italic=True, color=GREY)
h2('11.1 Daftar Fitur')
feat_table([
    ('Create Fit-out Project', 'Dari handover Fit-out → buat project + 1 task per item checklist.'),
    ('Tasks Smart Button', 'Akses task fit-out; field project_id pada handover.'),
])
h2('11.2 User Guide')
steps([
    'Buka handover bertipe Fit-out (dengan checklist terisi).',
    'Klik "Create Fit-out Project". Project & task otomatis terbuat.',
    'Klik "Fit-out Tasks" untuk mengelola pekerjaan di modul Project.',
])
h2('11.3 Skenario Test')
test_table([
    ('PJ-01', 'Buat project fit-out', 'Handover Fit-out dengan 6 item checklist → Create Fit-out Project.', 'project.project terbuat + 6 project.task; project_id terisi.'),
    ('PJ-02', 'Lihat tasks', 'Klik Fit-out Tasks.', 'Daftar 6 task pada project tampil.'),
])

# ---- DOCUMENTS ----
h1('12. Modul: Property Documents (Enterprise)')
para('Nama teknis: rental_management_documents — pada form Properti (tab Documents (App)).', italic=True, color=GREY)
h2('12.1 Daftar Fitur')
feat_table([
    ('Folder per Properti', 'documents.document folder per properti.'),
    ('Sync Attachments', 'Dorong ir.attachment properti ke Documents & buka folder.'),
])
h2('12.2 User Guide')
steps([
    'Pastikan app Documents (Enterprise) aktif.',
    'Buka properti → tab "Documents (App)" → klik "Open / Sync Documents".',
    'Folder dibuat (bila belum) dan lampiran properti tampil di Documents.',
])
h2('12.3 Skenario Test')
test_table([
    ('DM-01', 'Buat folder & sync', 'Properti dengan beberapa attachment → Open / Sync Documents.', 'Folder properti terbuat; dokumen muncul di folder; documents_count > 0.'),
])

# ============================ END-TO-END ============================
# ---- FIXED ASSET ----
h1('13. Modul: Fixed Assets, Depreciation & Revaluation')
para('Nama teknis: rental_management_asset — Menu: Properties → Fixed Assets.', italic=True, color=GREY)
h2('13.1 Daftar Fitur')
feat_table([
    ('Asset Register', 'Aset terhubung properti: nilai perolehan, salvage, metode, jumlah & periode penyusutan, akun & jurnal.'),
    ('Depreciation Board', 'Jadwal penyusutan Garis Lurus / Saldo Menurun; posting jurnal Dr Beban / Cr Akumulasi, ber-tag properti.'),
    ('Auto-post (Cron)', 'Cron harian memposting baris penyusutan yang jatuh tempo.'),
    ('Revaluation', 'Revaluasi naik/turun: posting ke Revaluation Reserve + recompute prospektif sisa board.'),
    ('Sync CORETAX L9', 'Buat entri register penyusutan CORETAX L9 dari aset (bila modul coretax terpasang).'),
])
h2('13.2 User Guide')
steps([
    'Properties → Fixed Assets → New: isi properti, tanggal & nilai perolehan, salvage, metode, jumlah/periode, akun & jurnal.',
    'Klik Compute Depreciation → tinjau board → Confirm (state Running).',
    'Klik Post Due Depreciation (atau biarkan cron harian).',
    'Untuk revaluasi: tambah baris Revaluation (+/-) → Post Revaluations (board sisa dihitung ulang atas nilai buku baru).',
    'Klik Sync CORETAX L9 untuk mendorong aset ke register penyusutan tahunan.',
])
h2('13.3 Skenario Test')
test_table([
    ('AS-01', 'Board garis lurus', 'Aset 10jt, salvage 0, linear, jumlah=5, periode=12. Compute Depreciation.', '5 baris @ 2jt; akumulasi akhir = 10jt; remaining akhir = 0.'),
    ('AS-02', 'Posting penyusutan', 'Confirm; set tanggal baris ≤ hari ini; Post Due Depreciation.', 'Jurnal Dr Beban/Cr Akumulasi terbuat & posted; ber-tag properti (analytic).'),
    ('AS-03', 'Saldo menurun', 'Aset declining, factor=2, jumlah=5. Compute.', 'Nominal menurun tiap periode; baris terakhir menutup ke salvage.'),
    ('AS-04', 'Revaluation naik', 'Tambah revaluation +5jt → Post Revaluations.', 'Jurnal Dr Aset/Cr Reserve 5jt; gross_value naik; sisa board dihitung ulang atas nilai buku baru.'),
    ('AS-05', 'Revaluation turun', 'Tambah revaluation −2jt → Post Revaluations.', 'Jurnal Dr Reserve/Cr Aset 2jt; gross_value turun; board ter-recompute.'),
    ('AS-06', 'Sync CORETAX L9', 'Klik Sync CORETAX L9 (modul coretax terpasang).', 'Entri coretax.asset.depreciation terbuat dengan nilai perolehan, remaining, penyusutan tahun berjalan.'),
])

# ---- DOCUMENT CE ----
h1('14. Modul: Property Documents (Community)')
para('Nama teknis: rental_management_document_ce — pada form Properti (tab Documents).', italic=True, color=GREY)
h2('14.1 Daftar Fitur')
feat_table([
    ('Register Lampiran', 'Daftar ir.attachment per properti (fallback Community jika app Documents Enterprise tidak tersedia).'),
    ('Kategori & Hitung', 'Kategori dokumen + smart button jumlah lampiran.'),
])
h2('14.2 User Guide')
steps([
    'Buka properti → tab "Documents" → unggah/lampirkan dokumen (kontrak, sertifikat, IMB).',
    'Pakai kategori untuk mengelompokkan; jumlah tampil pada smart button.',
])
h2('14.3 Skenario Test')
test_table([
    ('DC-01', 'Unggah lampiran', 'Lampirkan 2 dokumen pada properti via tab Documents.', 'documents_count=2; dokumen tersimpan sebagai ir.attachment milik properti.'),
])

# ---- OWNER PORTAL ----
h1('15. Modul: Owner Portal')
para('Nama teknis: rental_management_owner_portal — URL: /my/properties (portal pemilik).', italic=True, color=GREY)
h2('15.1 Daftar Fitur')
feat_table([
    ('/my/properties', 'Daftar properti milik pemilik yang login.'),
    ('Detail Properti', 'Ringkasan properti + daftar owner remittance terkait.'),
    ('Keamanan', 'Record rule membatasi tampilan hanya properti yang dimiliki user portal.'),
])
h2('15.2 User Guide')
steps([
    'Pastikan kontak pemilik tercantum sebagai owner pada properti dan diberi akses portal.',
    'Pemilik login → /my → kartu Properties → lihat daftar & detail properti.',
    'Buka detail untuk melihat remittance yang sudah diposting.',
])
h2('15.3 Skenario Test')
test_table([
    ('OP-01', 'Akses pemilik', 'Login sebagai owner portal yang memiliki properti.', 'Kartu/daftar Properties menampilkan hanya properti miliknya.'),
    ('OP-02', 'Isolasi data', 'Owner A coba akses properti owner B via id URL.', 'Akses ditolak / diarahkan ke /my.'),
    ('OP-03', 'Lihat remittance', 'Buka detail properti dengan remittance posted.', 'Daftar remittance tampil dengan jumlah benar.'),
])

# ---- CAM ----
h1('16. Modul: CAM / Service Charge')
para('Nama teknis: rental_management_cam — Menu: Properties → CAM / Service Charge.', italic=True, color=GREY)
h2('16.1 Daftar Fitur')
feat_table([
    ('Expense Pool', 'Kumpulan biaya CAM per properti/periode: budget vs actual.'),
    ('Apportionment', 'Alokasi ke tenant berdasarkan area share (m² / persentase).'),
    ('Invoice Service Charge', 'Tagih service charge ke tenant (ber-tag properti → laporan + analytic).'),
    ('Reconciliation', 'Selisih budget vs actual untuk penyesuaian akhir periode.'),
])
h2('16.2 User Guide')
steps([
    'Properties → CAM → New: pilih properti & periode; isi pool biaya (estimasi/budget).',
    'Isi area share tiap tenant (atau ambil dari luas unit).',
    'Klik Apportion → tinjau alokasi per tenant.',
    'Klik Create Invoices untuk menagih service charge.',
])
h2('16.3 Skenario Test')
test_table([
    ('CM-01', 'Apportion per area', 'Pool=120jt; 3 tenant area 100/200/300 m². Apportion.', 'Alokasi 20/40/60 jt sesuai proporsi area.'),
    ('CM-02', 'Invoice service charge', 'Klik Create Invoices.', 'out_invoice per tenant ber-link properti; muncul di Owners Statement & analytic.'),
    ('CM-03', 'Budget vs actual', 'Isi actual berbeda dari budget.', 'Selisih (variance) terhitung untuk rekonsiliasi.'),
])

# ---- RENT ESCALATION ----
h1('17. Modul: Rent Escalation')
para('Nama teknis: rental_management_rent_escalation — pada kontrak + cron.', italic=True, color=GREY)
h2('17.1 Daftar Fitur')
feat_table([
    ('Aturan Eskalasi', 'Kenaikan sewa berkala: fixed % atau fixed amount, interval (tahunan/periodik).'),
    ('Auto-apply (Cron)', 'Cron menerapkan kenaikan pada tanggal jatuh tempo eskalasi.'),
    ('Escalation Log', 'Riwayat kenaikan (tanggal, nilai lama → baru).'),
])
h2('17.2 User Guide')
steps([
    'Buka kontrak → bagian Rent Escalation: aktifkan, pilih tipe (%/amount), nilai & interval, tanggal mulai.',
    'Biarkan cron menerapkan kenaikan otomatis, atau jalankan manual untuk uji.',
    'Periksa Escalation Log untuk audit kenaikan.',
])
h2('17.3 Skenario Test')
test_table([
    ('RE-01', 'Eskalasi %', 'Sewa 10jt, escalation 10%/tahun, tanggal due hari ini. Jalankan cron.', 'Sewa naik ke 11jt; log mencatat 10jt→11jt.'),
    ('RE-02', 'Eskalasi amount', 'Escalation +500rb. Jalankan cron.', 'Sewa naik sebesar 500rb; log tercatat.'),
    ('RE-03', 'Belum jatuh tempo', 'Tanggal due di masa depan. Jalankan cron.', 'Tidak ada perubahan sewa.'),
])

# ---- DASHBOARD ----
h1('18. Modul: KPI Dashboard')
para('Nama teknis: rental_management_dashboard — Menu: Properties → Dashboard.', italic=True, color=GREY)
h2('18.1 Daftar Fitur')
feat_table([
    ('KPI Manajemen', 'Jumlah properti, kontrak aktif, NOI, arrears, collection rate, lease expiring 12 bulan.'),
    ('Filter', 'Per properti / periode.'),
])
h2('18.2 User Guide')
steps([
    'Properties → Dashboard: pilih periode (dan properti bila perlu).',
    'Tinjau kartu KPI; klik untuk drill-down ke daftar terkait bila tersedia.',
])
h2('18.3 Skenario Test')
test_table([
    ('DB-01', 'Hitung KPI', 'Dengan kontrak & invoice ada, buka Dashboard.', 'NOI = income − expense; arrears = invoice overdue; collection rate konsisten.'),
    ('DB-02', 'Lease expiring', 'Ada kontrak berakhir dalam 12 bulan.', 'Counter leases-expiring menampilkan jumlah benar.'),
])

# ---- DUNNING ----
h1('19. Modul: Dunning (Penagihan Otomatis)')
para('Nama teknis: rental_management_dunning — Menu: Accounting + cron.', italic=True, color=GREY)
h2('19.1 Daftar Fitur')
feat_table([
    ('Dunning Ladder', 'Tingkat penagihan bertahap untuk invoice jatuh tempo.'),
    ('Reminder Email', 'Kirim email template otomatis per level.'),
    ('Late Fee (opsional)', 'Tambah denda keterlambatan.'),
    ('Level Tracking', 'Dunning level tercatat per invoice (account.move).'),
])
h2('19.2 User Guide')
steps([
    'Konfigurasi level dunning (hari overdue, template email, late fee).',
    'Biarkan cron mengevaluasi invoice overdue & naikkan level + kirim reminder.',
    'Pantau dunning level pada invoice.',
])
h2('19.3 Skenario Test')
test_table([
    ('DN-01', 'Naik level', 'Invoice overdue melewati ambang level 1. Jalankan cron.', 'Dunning level=1; email reminder terkirim.'),
    ('DN-02', 'Late fee', 'Aktifkan late fee pada level. Jalankan cron.', 'Denda ditambahkan sesuai konfigurasi.'),
    ('DN-03', 'Lunas berhenti', 'Bayar invoice. Jalankan cron.', 'Tidak ada reminder lanjutan.'),
])

# ---- INSURANCE ----
h1('20. Modul: Insurance Register')
para('Nama teknis: rental_management_insurance — Menu: Properties → Insurance Policies.', italic=True, color=GREY)
h2('20.1 Daftar Fitur')
feat_table([
    ('Polis Asuransi', 'Register polis properti/gedung: penanggung, nomor, nilai, periode.'),
    ('Alert Expiry (Cron)', 'Reminder activity menjelang kedaluwarsa polis.'),
])
h2('20.2 User Guide')
steps([
    'Properties → Insurance Policies → New: pilih properti, penanggung, nomor, nilai pertanggungan, periode, lead reminder.',
    'Pantau filter Expiring Soon; perpanjang sebelum kedaluwarsa.',
])
h2('20.3 Skenario Test')
test_table([
    ('IN-01', 'Buat polis', 'Buat polis dengan expiry 30 hari ke depan.', 'Polis tersimpan; days-to-expiry terhitung.'),
    ('IN-02', 'Reminder', 'Expiry dalam lead window. Jalankan cron.', 'Activity reminder terjadwal pada penanggung jawab.'),
])

# ---- LEASE EXPIRY ----
h1('21. Modul: Lease Expiry & Renewal')
para('Nama teknis: rental_management_lease_expiry — Menu: Properties → Lease Expiry + cron.', italic=True, color=GREY)
h2('21.1 Daftar Fitur')
feat_table([
    ('Daftar & Filter', 'Kontrak yang berakhir dalam 30 / 90 hari.'),
    ('Reminder Renewal (Cron)', 'Activity reminder perpanjangan otomatis.'),
])
h2('21.2 User Guide')
steps([
    'Properties → Lease Expiry: gunakan filter 30/90 hari untuk melihat kontrak mendekati akhir.',
    'Biarkan cron menjadwalkan reminder renewal; tindak lanjuti perpanjangan.',
])
h2('21.3 Skenario Test')
test_table([
    ('LE-01', 'Filter 30 hari', 'Kontrak berakhir dalam 20 hari.', 'Muncul pada filter 30 hari.'),
    ('LE-02', 'Reminder renewal', 'Jalankan cron.', 'Activity reminder renewal terjadwal.'),
])

# ---- VACANCY ----
h1('22. Modul: Vacancy / Availability Board')
para('Nama teknis: rental_management_vacancy — Menu: Properties → Vacancy.', italic=True, color=GREY)
h2('22.1 Daftar Fitur')
feat_table([
    ('Board Status', 'Properti dikelompokkan per status: available / on-rent / booked.'),
    ('Filter', 'Penyaringan cepat unit kosong vs tersewa.'),
])
h2('22.2 User Guide')
steps([
    'Properties → Vacancy: lihat pengelompokan status.',
    'Filter "Available" untuk daftar unit yang siap dipasarkan.',
])
h2('22.3 Skenario Test')
test_table([
    ('VC-01', 'Kelompok status', 'Beberapa properti dengan status berbeda.', 'Board mengelompokkan benar sesuai stage.'),
    ('VC-02', 'Filter available', 'Pilih filter Available.', 'Hanya properti available tampil.'),
])

# ---- PPM ----
h1('23. Modul: Preventive Maintenance (PPM)')
para('Nama teknis: rental_management_ppm — Menu: Properties/Maintenance + cron.', italic=True, color=GREY)
h2('23.1 Daftar Fitur')
feat_table([
    ('PPM Plan', 'Jadwal preventive maintenance per properti/peralatan dengan interval & SLA.'),
    ('Auto-generate (Cron)', 'Membuat maintenance.request otomatis sesuai jadwal.'),
])
h2('23.2 User Guide')
steps([
    'Buat PPM Plan: pilih properti/peralatan, interval (mis. bulanan), SLA, penanggung jawab.',
    'Biarkan cron membuat maintenance.request pada jadwal; kerjakan via modul Maintenance.',
])
h2('23.3 Skenario Test')
test_table([
    ('PP-01', 'Generate request', 'PPM plan due hari ini. Jalankan cron.', 'maintenance.request baru terbuat sesuai plan.'),
    ('PP-02', 'Interval berikut', 'Setelah generate, cek next date.', 'Tanggal jadwal berikut bergeser sesuai interval.'),
])

# ---- PARKING ----
h1('24. Modul: Parking Management')
para('Nama teknis: rental_management_parking — Menu: Properties → Parking.', italic=True, color=GREY)
h2('24.1 Daftar Fitur')
feat_table([
    ('Bay Register', 'Daftar petak parkir per properti dengan status.'),
    ('Alokasi Tenant', 'Tetapkan bay ke tenant/kontrak.'),
    ('Invoice Parkir', 'Tagih sewa parkir (ber-tag properti → laporan + analytic).'),
])
h2('24.2 User Guide')
steps([
    'Properties → Parking: daftarkan bay (nomor, tipe, tarif).',
    'Alokasikan bay ke tenant; buat invoice sewa parkir.',
])
h2('24.3 Skenario Test')
test_table([
    ('PK-01', 'Alokasi bay', 'Tetapkan bay ke tenant.', 'Bay berstatus allocated; terkait tenant/kontrak.'),
    ('PK-02', 'Invoice parkir', 'Buat invoice parkir.', 'out_invoice ber-link properti; muncul di Owners Statement.'),
])

# ---- VALUATION ----
h1('25. Modul: Property Valuation')
para('Nama teknis: rental_management_valuation — Menu: Properties → Valuations.', italic=True, color=GREY)
h2('25.1 Daftar Fitur')
feat_table([
    ('Register Valuasi', 'Valuasi pasar berkala per properti (tanggal, penilai, metode, nilai).'),
    ('Latest Value', 'Nilai terkini tampil pada properti.'),
])
h2('25.2 User Guide')
steps([
    'Properties → Valuations → New: pilih properti, tanggal, penilai, metode, nilai pasar.',
    'Nilai terbaru otomatis menjadi latest value pada properti.',
])
h2('25.3 Skenario Test')
test_table([
    ('VL-01', 'Catat valuasi', 'Buat 2 valuasi tanggal berbeda.', 'Latest value = valuasi tanggal terbaru.'),
])

# ---- ACCESS ----
h1('26. Modul: Access Cards & Visitor Log')
para('Nama teknis: rental_management_access — Menu: Properties → Access.', italic=True, color=GREY)
h2('26.1 Daftar Fitur')
feat_table([
    ('Access/Parking Card', 'Register kartu akses/parkir per tenant.'),
    ('Visitor Log', 'Log check-in/check-out tamu.'),
])
h2('26.2 User Guide')
steps([
    'Properties → Access: terbitkan kartu (nomor, tenant, masa berlaku).',
    'Catat kunjungan tamu pada Visitor Log (check-in/out).',
])
h2('26.3 Skenario Test')
test_table([
    ('AC-01', 'Terbitkan kartu', 'Buat kartu akses untuk tenant.', 'Kartu tercatat & terkait tenant.'),
    ('AC-02', 'Log tamu', 'Catat check-in lalu check-out tamu.', 'Waktu check-in/out tersimpan.'),
])

# ---- ESIGN ----
h1('27. Modul: Contract e-Signature Tracking')
para('Nama teknis: rental_management_esign — pada kontrak + mail.', italic=True, color=GREY)
h2('27.1 Daftar Fitur')
feat_table([
    ('Signature Request', 'Tracking permintaan tanda tangan kontrak (ringan, tanpa app Sign Enterprise).'),
    ('Status', 'Sent → Signed; catatan tanggal & penandatangan via email/chatter.'),
])
h2('27.2 User Guide')
steps([
    'Buka kontrak → Request Signature: kirim permintaan ke penandatangan.',
    'Tandai Signed saat dokumen ditandatangani; status & tanggal tercatat.',
])
h2('27.3 Skenario Test')
test_table([
    ('ES-01', 'Kirim permintaan', 'Klik Request Signature pada kontrak.', 'Status=Sent; nomor urut sequence terbuat; email/log tercatat.'),
    ('ES-02', 'Tandai signed', 'Klik Mark Signed.', 'Status=Signed; tanggal tanda tangan tersimpan.'),
])

h1('28. Skenario Test End-to-End (Alur Lengkap)')
para('Skenario terintegrasi lintas modul untuk memvalidasi alur bisnis penuh.')
test_table([
    ('E2E-1', 'Onboarding tenant', 'Lead (CRM) → Create Lease Contract → aktifkan kontrak → Handover Move-in → Guarantee dicatat.', 'Kontrak aktif; handover selesai; jaminan Active; tenant bisa login portal.'),
    ('E2E-2', 'Billing bulanan', 'Invoice sewa + Meter recharge + GTO turnover diposting.', 'Semua pendapatan ber-link properti & analytic; muncul di Owners Statement.'),
    ('E2E-3', 'Pengeluaran', 'Maintenance → Create PO → terima → vendor bill posted.', 'Beban masuk Payment Details & analytic properti.'),
    ('E2E-4', 'Tutup periode', 'Cetak Owners Statement; buat Owner Remittance; ekspor CORETAX Faktur PK.', 'Laporan konsisten; remittance terposting; XML e-Faktur terunduh.'),
    ('E2E-5', 'Move-out', 'Handover Move-out + make-good; Security Deposit refund/deduction; kontrak ditutup.', 'Deposit terselesaikan; Sec Dep Bal nol; kontrak closed.'),
])

para(' ')
para('Catatan akhir: Seluruh modul telah lolos validasi sintaks Python/XML dan struktur, '
     'namun WAJIB dilakukan smoke test pada instance Odoo 19 (staging) sebelum produksi, '
     'termasuk validasi XML CORETAX terhadap XSD resmi DJP dan konfirmasi hook tampilan '
     'antar versi minor Odoo.', italic=True, color=GREY, size=9.5)

doc.save('../../docs/Custom_Modules_Feature_List_UserGuide_TestScenarios.docx')
print('SAVED OK')
