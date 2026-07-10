# -*- coding: utf-8 -*-
"""User Guide & Deskripsi Fitur — disusun mengikuti alur bisnis, bahasa awam."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x39, 0x64)
GREEN = RGBColor(0x00, 0x6A, 0x4E)
GREY = RGBColor(0x55, 0x55, 0x55)
AMBER = RGBColor(0xB4, 0x6A, 0x00)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)


def shade(c, h):
    tcPr = c._tc.get_or_add_tcPr(); s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), h)
    tcPr.append(s)


def ct(c, t, bold=False, white=False, size=9, color=None):
    c.text = ''; p = c.paragraphs[0]; r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color: r.font.color.rgb = color


def bab(no, t):
    p = doc.add_heading(level=1)
    r = p.add_run('BAB %s — %s' % (no, t)); r.font.color.rgb = NAVY; r.font.size = Pt(16)
    return p


def h2(t):
    p = doc.add_heading(level=2); r = p.add_run(t); r.font.color.rgb = GREEN; r.font.size = Pt(12.5); return p


def h3(t):
    p = doc.add_heading(level=3); r = p.add_run(t); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33); r.font.size = Pt(11)
    return p


def para(t, bold=False, italic=False, size=10.5, color=None, after=6):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p


def bullet(t, pre=None, size=10):
    p = doc.add_paragraph(style='List Bullet')
    if pre:
        rr = p.add_run(pre); rr.bold = True; rr.font.size = Pt(size)
    r = p.add_run(t); r.font.size = Pt(size)
    return p


def steps(items):
    for s in items:
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(s); r.font.size = Pt(10)


def tip(t):
    p = doc.add_paragraph()
    r1 = p.add_run('Tips: '); r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = AMBER
    r2 = p.add_run(t); r2.italic = True; r2.font.size = Pt(9.5); r2.font.color.rgb = AMBER
    p.paragraph_format.space_after = Pt(10)


def fitur(nama, teknis, deskripsi, manfaat, langkah, tips=None):
    """Blok deskripsi 1 fitur/modul: nama, nama teknis modul, deskripsi awam,
    manfaat bisnis, langkah pakai, tips opsional."""
    h3(nama)
    p = doc.add_paragraph(); r = p.add_run('Modul teknis: %s' % teknis)
    r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(4)
    para(deskripsi, after=4)
    pm = doc.add_paragraph(); rm = pm.add_run('Manfaat: '); rm.bold = True; rm.font.size = Pt(10)
    rm2 = pm.add_run(manfaat); rm2.font.size = Pt(10)
    pm.paragraph_format.space_after = Pt(4)
    pl = doc.add_paragraph(); rl = pl.add_run('Cara menggunakan:'); rl.bold = True; rl.font.size = Pt(10)
    pl.paragraph_format.space_after = Pt(2)
    steps(langkah)
    if tips:
        tip(tips)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def toc_field():
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-2" \\h \\z \\u')
    run._r.addnext(fld)
    t = OxmlElement('w:t'); t.text = '(Klik kanan → Update Field untuk memuat Daftar Isi)'
    rpr = OxmlElement('w:r'); rpr.append(t)
    fld.append(rpr)


# ==================== COVER ====================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('PANDUAN PENGGUNA & DESKRIPSI FITUR')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Property Management System — Odoo 19\nDisusun mengikuti alur proses bisnis pengelolaan properti, dari awal hingga akhir')
r.font.size = Pt(13); r.bold = True; r.font.color.rgb = GREEN
doc.add_paragraph()
meta = doc.add_table(rows=6, cols=2); meta.style = 'Table Grid'
for i, (k, v) in enumerate([
    ('Dokumen', 'User Guide & Feature Description (bahasa non-teknis)'),
    ('Platform', 'Odoo 19 Community + addon rental_management + 26 modul companion'),
    ('Ditujukan untuk', 'User Operasional, Accounting/Finance, Manajemen, Property Manager'),
    ('Cara membaca', 'Ikuti urutan BAB sesuai alur bisnis: dari master data → leasing → operasional → tutup periode'),
    ('Disusun oleh', 'System Analyst / Functional & Odoo Developer'),
    ('Tanggal', '9 Juli 2026'),
]):
    shade(meta.rows[i].cells[0], 'E7ECF5'); ct(meta.rows[i].cells[0], k, bold=True, size=10)
    ct(meta.rows[i].cells[1], v, size=10)
    meta.rows[i].cells[0].width = Inches(1.8); meta.rows[i].cells[1].width = Inches(4.9)
doc.add_page_break()

para('DAFTAR ISI', bold=True, size=13, color=NAVY, after=6)
toc_field()
doc.add_page_break()

# ==================== BAB 1 — PENDAHULUAN ====================
bab(1, 'Pendahuluan')
h2('1.1 Tentang Sistem Ini')
para('Sistem ini adalah aplikasi pengelolaan properti terpadu yang dibangun di atas Odoo 19. '
     'Satu sistem ini menangani seluruh siklus hidup properti — mulai dari mencari tenant, membuat '
     'kontrak sewa, menagih bulanan, merawat gedung, sampai menyusun laporan keuangan untuk pemilik '
     'dan melaporkan pajak. Semua data saling terhubung: begitu sebuah properti, kontrak, atau tagihan '
     'dibuat, seluruh modul lain otomatis "mengenali" dan menggunakannya — tidak perlu input ulang.')
h2('1.2 Siapa yang Menggunakan Sistem Ini')
table = doc.add_table(rows=1, cols=2); table.style = 'Table Grid'
for i, htxt in enumerate(['Peran', 'Aktivitas Utama di Sistem']):
    shade(table.rows[0].cells[i], '1F3964'); ct(table.rows[0].cells[i], htxt, bold=True, white=True, size=9.5)
for role, act in [
    ('Tim Operasional / Leasing', 'Kelola prospek tenant, buat kontrak, handover unit, kelola parkir & akses, tindak lanjuti perpanjangan sewa.'),
    ('Tim Accounting / Finance', 'Kelola tagihan, terima pembayaran, susun anggaran, cetak laporan pemilik, kelola pajak (CORETAX), kelola aset.'),
    ('Property Manager', 'Pantau okupansi, kelola pemeliharaan gedung, koordinasi dengan tenant & vendor.'),
    ('Manajemen', 'Memantau dashboard KPI (pendapatan, tunggakan, okupansi) untuk pengambilan keputusan.'),
    ('Pemilik Properti (Owner)', 'Melihat laporan keuangan propertinya & riwayat penyaluran dana lewat Owner Portal.'),
    ('Tenant (Penyewa)', 'Melihat kontrak, tagihan, dan mengajukan permintaan lewat Tenant Portal.'),
]:
    cells = table.add_row().cells
    ct(cells[0], role, bold=True, size=9); ct(cells[1], act, size=9)
    cells[0].width = Inches(1.8); cells[1].width = Inches(4.9)
doc.add_paragraph()
h2('1.3 Alur Besar yang Dipakai dalam Panduan Ini')
para('Panduan ini TIDAK disusun berdasarkan daftar modul teknis, melainkan mengikuti tahapan nyata '
     'yang dijalani sebuah properti dan tenant, agar mudah diikuti:')
steps([
    'Menyiapkan data dasar (properti, unit, pemilik, produk).',
    'Menjaring calon tenant & membuat kontrak sewa.',
    'Menyerahkan unit ke tenant (onboarding / move-in).',
    'Operasional harian: menagih, menerima pembayaran, mengingatkan tunggakan.',
    'Merawat gedung & mengadakan barang/jasa.',
    'Menutup periode: laporan keuangan pemilik, anggaran, pajak, aset.',
    'Layanan mandiri untuk tenant & pemilik lewat portal.',
    'Mengelola risiko: jaminan, asuransi, valuasi.',
    'Akhir sewa: perpanjangan atau tenant keluar (move-out).',
    'Memantau semuanya lewat dashboard KPI.',
])
doc.add_page_break()

# ==================== BAB 2 — MASTER DATA ====================
bab(2, 'Menyiapkan Data Dasar')
para('Sebelum sistem bisa dipakai sehari-hari, beberapa data dasar perlu disiapkan lebih dulu. '
     'Data ini menjadi "pondasi" — dipakai berulang oleh hampir semua fitur lain.')

fitur('Properti & Unit', 'rental_management (dasar)',
      'Tempat mencatat seluruh gedung/aset yang dikelola: nama properti, alamat, jenis (kantor, ritel, '
      'apartemen, dsb), dan daftar unit/lantai di dalamnya beserta luasnya.',
      'Menjadi acuan tunggal seluruh transaksi — kontrak, tagihan, laporan keuangan, dan pemeliharaan '
      'semuanya merujuk ke data properti ini.',
      ['Buka menu Properties → Properties → klik New.',
       'Isi nama, alamat, jenis properti, dan tambahkan unit-unit di dalamnya (nomor unit, luas, status).',
       'Simpan. Properti siap dipakai untuk membuat kontrak.'])

fitur('Kepemilikan Properti', 'rental_management_financial_report',
      'Mencatat siapa pemilik sebuah properti. Satu properti bisa dimiliki lebih dari satu pihak, '
      'masing-masing dengan persentase kepemilikannya.',
      'Menjadi dasar pembagian hasil sewa (remittance) ke tiap pemilik secara proporsional, dan '
      'menentukan siapa yang berhak melihat laporan properti tersebut lewat Owner Portal.',
      ['Buka properti → tab "Owners & Financial".',
       'Tambahkan baris pemilik beserta persentase kepemilikan (total harus 100%).',
       'Isi juga Property Manager yang bertanggung jawab atas properti ini.'],
      tips='Sisihkan waktu untuk memastikan total persentase kepemilikan selalu 100% — bila tidak, '
           'perhitungan pembagian dana ke pemilik akan salah.')

fitur('Produk & Layanan', 'product (standar Odoo, dipakai semua modul)',
      'Daftar "barang/jasa" yang ditagihkan ke tenant: sewa, service charge, recharge listrik, parkir, '
      'dan lain-lain. Tiap produk terhubung ke akun pendapatan tertentu di pembukuan.',
      'Membuat penagihan konsisten dan otomatis masuk ke kategori laporan yang benar.',
      ['Buka Accounting/Sales → Products.',
       'Pastikan tiap produk (Sewa, Service Charge, Utility Recharge, Parking, dst.) punya akun pendapatan yang benar.',
       'Produk baru otomatis tersedia saat menagih tenant di modul manapun.'])
doc.add_page_break()

# ==================== BAB 3 — LEASING PIPELINE ====================
bab(3, 'Menjaring Calon Tenant')
fitur('Pipeline Leasing (CRM)', 'rental_management_crm',
      'Corong (funnel) untuk mengelola calon tenant sejak pertama kali tertarik hingga sepakat menyewa — '
      'menyerupai CRM penjualan pada umumnya, namun sudah disesuaikan untuk leasing properti.',
      'Tim leasing tidak lagi kehilangan jejak prospek; setiap tahap (tertarik → survei → negosiasi → '
      'deal) tercatat rapi, dan begitu deal, kontrak bisa langsung dibuat tanpa input ulang data.',
      ['Buka menu CRM → buat kartu prospek baru, isi unit yang diminati, target sewa, dan durasi.',
       'Geser kartu prospek melintasi tahapan (kolom) sesuai progresnya.',
       'Setelah disepakati, klik tombol "Create Lease Contract" pada kartu prospek.',
       'Form kontrak sewa akan terbuka otomatis dengan data unit & tenant sudah terisi — lengkapi sisanya lalu simpan.'],
      tips='Gunakan smart button pada kartu kontak tenant untuk melihat riwayat semua kontrak yang '
           'pernah/sedang berjalan dengan tenant tersebut.')
doc.add_page_break()

# ==================== BAB 4 — KONTRAK SEWA ====================
bab(4, 'Membuat & Mengelola Kontrak Sewa')
fitur('Kontrak Sewa Standar', 'rental_management (dasar)',
      'Dokumen inti yang mengikat tenant dengan unit tertentu: masa sewa, nilai sewa, jadwal penagihan, '
      'dan syarat lainnya.',
      'Menjadi sumber kebenaran tunggal untuk seluruh penagihan otomatis selama masa sewa berjalan.',
      ['Buat kontrak baru (langsung, atau otomatis dari CRM).',
       'Pilih unit, tenant, tanggal mulai & akhir sewa, serta nilai sewa dan jadwal tagihan.',
       'Aktifkan kontrak setelah semua syarat (jaminan, deposit) terpenuhi.'])

fitur('Sewa Jangka Pendek (Casual Leasing)', 'rental_management_casual_leasing',
      'Untuk sewa singkat yang bukan kontrak jangka panjang — misalnya booth event, kios pop-up, atau '
      'area promosi harian/mingguan.',
      'Memudahkan penagihan sewa jangka pendek tanpa harus membuat kontrak formal jangka panjang.',
      ['Buka Casual Leases → New: pilih properti/space, tenant, periode sewa, dan tarif (harian/mingguan/tetap).',
       'Sistem otomatis menghitung total tagihan berdasarkan durasi & tarif.',
       'Klik Confirm, lalu Create Invoice untuk menagih.'])

fitur('Jaminan Bank / Asuransi (Guarantee)', 'rental_management_guarantee',
      'Pencatatan jaminan yang diserahkan tenant sebagai syarat sewa — bisa berupa bank guarantee atau '
      'jaminan asuransi.',
      'Sistem otomatis mengingatkan sebelum jaminan kedaluwarsa, sehingga tim tidak lupa memperpanjang '
      'atau mencairkan sesuai kebutuhan.',
      ['Catat jaminan pada kontrak: jenis, penerbit, nomor, nilai, tanggal terbit & kedaluwarsa.',
       'Aktifkan jaminan setelah dokumen fisik diterima dan diverifikasi.',
       'Sistem mengirim pengingat otomatis menjelang tanggal kedaluwarsa.'])

fitur('Uang Jaminan / Security Deposit', 'rental_management_financial_report',
      'Deposit yang dibayar tenant di awal sewa sebagai jaminan. Deposit ini dicatat sebagai kewajiban '
      '(bukan pendapatan), sehingga tidak salah masuk sebagai laba.',
      'Saldo deposit tenant selalu terlihat jelas, dan pemotongan/pengembaliannya tercatat rapi saat '
      'tenant keluar.',
      ['Catat penerimaan deposit pada kontrak → tandai "Held" setelah diterima.',
       'Saat tenant keluar: catat pemotongan (bila ada kerusakan/tunggakan) sebelum mengembalikan sisa.',
       'Sisa saldo deposit selalu terlihat pada laporan Tenant Balances.'])

fitur('Kenaikan Sewa Berkala (Rent Escalation)', 'rental_management_rent_escalation',
      'Aturan kenaikan sewa otomatis pada tanggal tertentu (misalnya tiap ulang tahun kontrak), baik '
      'dalam bentuk persentase maupun nominal tetap.',
      'Kenaikan sewa tidak lagi perlu diingat & diinput manual setiap tahun — sistem menjalankannya '
      'sendiri sesuai aturan yang sudah diatur.',
      ['Pada kontrak, aktifkan Rent Escalation: pilih tipe (persentase/nominal), nilai, dan tanggal berlaku.',
       'Sistem otomatis menaikkan sewa pada tanggal yang ditentukan dan mencatatnya di Escalation Log.'])
doc.add_page_break()

# ==================== BAB 5 — ONBOARDING TENANT ====================
bab(5, 'Menyerahkan Unit ke Tenant (Move-in)')
fitur('Handover Move-in / Fit-out / Move-out', 'rental_management_handover (+ project)',
      'Checklist terstruktur untuk proses serah-terima unit — baik saat tenant masuk (move-in), masa '
      'renovasi tenant (fit-out), maupun saat tenant keluar (move-out).',
      'Tidak ada langkah yang terlewat saat serah-terima; kondisi unit terdokumentasi dengan jelas '
      'sehingga tidak terjadi perselisihan di kemudian hari.',
      ['Buat Handover baru, pilih tipe (Move-in/Fit-out/Move-out) dan kontrak terkait.',
       'Klik "Load Default Checklist" untuk memuat daftar tugas standar sesuai tipe.',
       'Centang tiap item saat selesai — progress bar akan menunjukkan persentase penyelesaian.',
       'Lampirkan foto/dokumen kondisi unit lewat kolom percakapan (chatter).',
       'Klik Complete saat seluruh proses selesai.'],
      tips='Untuk proses Fit-out yang kompleks, gunakan tombol "Create Fit-out Project" agar setiap item '
           'checklist berubah menjadi tugas (task) yang bisa dipantau progresnya di modul Project.')

fitur('Kartu Akses & Buku Tamu', 'rental_management_access',
      'Pencatatan kartu akses/parkir yang diberikan ke tenant, serta log kunjungan tamu ke gedung.',
      'Memudahkan kontrol keamanan gedung — siapa yang punya akses, dan siapa saja yang berkunjung.',
      ['Terbitkan kartu akses baru untuk tenant: nomor kartu & masa berlaku.',
       'Untuk tamu, catat waktu check-in dan check-out pada Visitor Log.'])

fitur('Alokasi Tempat Parkir', 'rental_management_parking',
      'Pengelolaan bay/petak parkir gedung — mana yang tersedia, mana yang sudah dialokasikan ke tenant.',
      'Parkir dikelola dan ditagih secara rapi, tidak lagi manual di luar sistem.',
      ['Daftarkan bay parkir yang ada beserta tarifnya.',
       'Alokasikan bay ke tenant/kontrak tertentu.',
       'Buat invoice sewa parkir secara berkala seperti tagihan sewa lainnya.'])
doc.add_page_break()

# ==================== BAB 6 — OPERASIONAL & PENAGIHAN ====================
bab(6, 'Operasional Harian & Penagihan')
para('Setelah tenant aktif menyewa, inilah rutinitas bulanan yang paling sering dilakukan tim operasional & finance.')

fitur('Sewa Berbasis Omzet (GTO / Revenue Sharing)', 'rental_management_gto_meter',
      'Untuk tenant yang membayar sewa berdasarkan omzet penjualannya (umum di ritel/mal) — bukan '
      'nilai tetap. Ada tiga model: ambil yang lebih besar antara sewa dasar atau persentase omzet, '
      'sewa dasar + kelebihan di atas ambang tertentu, atau murni persentase omzet.',
      'Pendapatan sewa menyesuaikan performa bisnis tenant secara otomatis, tanpa kalkulasi manual '
      'yang rawan salah.',
      ['Pada kontrak, aktifkan GTO dan pilih model perhitungannya serta persentasenya.',
       'Setiap periode, catat deklarasi omzet tenant pada menu GTO Turnovers.',
       'Sistem menghitung otomatis bagian yang harus ditagih, lalu klik Create Invoice.'])

fitur('Meter Listrik/Air & Recharge Utilitas', 'rental_management_gto_meter',
      'Pencatatan meteran listrik/air/gas per unit dan penagihan pemakaian ke tenant (recharge).',
      'Biaya utilitas tertagih akurat sesuai pemakaian aktual, bukan estimasi.',
      ['Daftarkan meter untuk tiap unit beserta tarif per satuannya.',
       'Setiap periode, input pembacaan meter terbaru (pembacaan sebelumnya otomatis terisi).',
       'Klik Create Recharge Invoice — sistem menghitung konsumsi dan nilai tagihan otomatis.'])

fitur('Biaya Pengelolaan Bersama (CAM / Service Charge)', 'rental_management_cam',
      'Pengumpulan biaya pengelolaan area bersama (kebersihan, keamanan, listrik area umum, dsb.) '
      'yang kemudian dibagi ke tenant sesuai proporsi luas unitnya.',
      'Biaya bersama dibagi secara adil dan transparan, dengan kemungkinan rekonsiliasi bila ada '
      'selisih antara anggaran dan realisasi.',
      ['Buat pool biaya CAM per periode, masukkan estimasi/anggaran biaya.',
       'Sistem menghitung porsi tiap tenant berdasarkan luas unit dibanding total luas gedung.',
       'Klik Create Invoices untuk menagih service charge ke seluruh tenant sekaligus.'])

fitur('Penagihan Otomatis & Pengingat Tunggakan (Dunning)', 'rental_management_dunning',
      'Sistem penagihan bertingkat otomatis untuk invoice yang menunggak — mulai dari pengingat halus '
      'hingga peringatan lebih tegas, bisa disertai denda keterlambatan.',
      'Tim finance tidak perlu memantau manual satu-per-satu; tunggakan ditindaklanjuti otomatis '
      'dan konsisten sesuai aturan yang ditetapkan.',
      ['Atur tingkatan penagihan (misalnya: pengingat di hari ke-7, peringatan di hari ke-14, dst.) '
       'beserta template emailnya.',
       'Sistem otomatis mengevaluasi invoice terlambat setiap hari dan mengirim pengingat sesuai levelnya.',
       'Bila diaktifkan, denda keterlambatan otomatis ditambahkan ke tagihan.'])
doc.add_page_break()

# ==================== BAB 7 — PEMELIHARAAN & PENGADAAN ====================
bab(7, 'Merawat Gedung & Pengadaan Barang/Jasa')
fitur('Permintaan Perbaikan (Maintenance)', 'maintenance (standar Odoo, diperluas rental_management)',
      'Pencatatan permintaan perbaikan/perawatan pada properti — baik yang dilaporkan tenant maupun '
      'temuan tim internal.',
      'Semua permintaan perbaikan tercatat, terlacak progresnya, dan biayanya otomatis terhubung ke '
      'properti terkait untuk keperluan laporan keuangan.',
      ['Buat permintaan maintenance baru: pilih properti, jenis masalah, dan tingkat urgensi.',
       'Tugaskan ke teknisi/vendor terkait, pantau statusnya hingga selesai.',
       'Bila perlu barang/jasa dari luar, klik "Create Purchase Order" langsung dari permintaan ini.'])

fitur('Perawatan Terjadwal (Preventive Maintenance / PPM)', 'rental_management_ppm',
      'Jadwal perawatan rutin untuk peralatan gedung (lift, genset, AHU, sistem kebakaran, dll.) yang '
      'dibuat otomatis sesuai frekuensi yang ditetapkan — tanpa menunggu rusak dulu.',
      'Peralatan gedung terawat secara proaktif, mengurangi risiko kerusakan mendadak dan biaya darurat.',
      ['Buat rencana PPM: pilih peralatan/area, frekuensi perawatan (misalnya bulanan), dan SLA penyelesaian.',
       'Sistem otomatis membuat permintaan maintenance baru setiap kali jadwal jatuh tempo.'])

fitur('Pengadaan Barang/Jasa (Purchase)', 'rental_management_purchase',
      'Proses pembelian barang/jasa untuk keperluan properti — mulai dari permintaan pembelian hingga '
      'tagihan vendor — yang terhubung langsung ke properti/kontrak/permintaan maintenance terkait.',
      'Setiap pengeluaran otomatis tercatat ke properti yang benar, sehingga muncul akurat di laporan '
      'keuangan pemilik tanpa perlu pencatatan manual ulang.',
      ['Buat Purchase Order, isi tab "Property" untuk mengaitkannya ke properti/kontrak/maintenance.',
       'Konfirmasi PO, terima barang/jasa, lalu klik Create Bill.',
       'Tagihan vendor otomatis mewarisi keterkaitan properti dari PO.'])
doc.add_page_break()

# ==================== BAB 8 — TUTUP PERIODE & KEUANGAN ====================
bab(8, 'Menutup Periode & Laporan Keuangan Pemilik')
para('Inilah inti dari sistem ini: menghasilkan laporan keuangan bulanan untuk pemilik properti, '
     'selayaknya laporan yang biasa diberikan oleh perusahaan pengelola properti profesional (gaya '
     'CBRE/MRI Owners Statement).')

fitur('Pemetaan Akun ke Kategori Laporan', 'rental_management_financial_report',
      'Setiap akun pendapatan/beban di pembukuan dipetakan ke kategori tertentu (misalnya "Rental '
      'Income", "Statutory Outgoings") agar laporan pemilik tersusun rapi per kelompok, bukan sekadar '
      'daftar akun mentah.',
      'Laporan pemilik tersaji dalam format yang informatif dan mudah dibaca, bukan tabel akuntansi mentah.',
      ['Buka Chart of Accounts, buka akun terkait, isi field "Property Report Category".',
       'Lakukan sekali di awal implementasi untuk seluruh akun income/expense terkait properti.'])

fitur('Anggaran Properti (Budget)', 'rental_management_financial_report',
      'Penyusunan anggaran pendapatan/beban per properti per periode, untuk kemudian dibandingkan '
      'dengan realisasi (Actual vs Budget vs Variance).',
      'Manajemen & pemilik dapat melihat seberapa baik performa aktual properti dibanding target '
      'yang direncanakan.',
      ['Buat Property Budget baru, isi target per akun per bulan.',
       'Ajukan untuk disetujui (Submit for Approval) sesuai jenjang otorisasi yang berlaku.',
       'Setelah disetujui, angka budget otomatis dibandingkan dengan realisasi pada Owners Statement.'])

fitur('Penampungan & Penyaluran Dana ke Pemilik (Trust & Remittance)', 'rental_management_financial_report',
      'Pemantauan saldo dana milik pemilik yang terkumpul dari sewa, dan proses penyalurannya kembali '
      'ke pemilik sesuai porsi kepemilikan masing-masing.',
      'Penyaluran dana ke pemilik tercatat rapi dan otomatis terbagi sesuai persentase kepemilikan — '
      'tidak perlu hitung manual satu per satu.',
      ['Buka menu Owner Remittances → New, pilih properti & tanggal.',
       'Klik "Compute from Owners" — sistem menghitung otomatis pembagian sesuai persentase kepemilikan.',
       'Periksa hasilnya, lalu klik Post untuk mencatatnya secara resmi.'])

fitur('Owners Statement — Laporan Lengkap untuk Pemilik', 'rental_management_financial_report',
      'Satu paket laporan bulanan berisi 9 bagian: Ringkasan Performa, Rincian Pendapatan & Beban, '
      'Penerimaan & Pembayaran Kas, Saldo Tenant, Tunggakan, Detail Pembayaran, Neraca Saldo, Neraca, '
      'dan Rekonsiliasi PPN.',
      'Pemilik menerima laporan yang komprehensif dan profesional setiap bulan, tanpa tim finance '
      'harus menyusunnya manual dari berbagai sumber data.',
      ['Buka menu Owners Statement, pilih properti dan periode yang ingin dicetak.',
       'Klik "Print Owners Statement" — sistem otomatis menarik data dari seluruh transaksi bulan itu.',
       'Bagikan PDF ke pemilik, atau arahkan pemilik untuk melihatnya sendiri lewat Owner Portal.'])

fitur('Pencatatan Otomatis per Properti (Analytic)', 'rental_management_financial_report',
      'Fitur "di balik layar" yang membuat setiap transaksi (pendapatan maupun beban) secara otomatis '
      'tercatat ke properti asalnya, tanpa perlu diinput manual setiap kali.',
      'Menjamin semua modul (sewa, GTO, meter, CAM, parkir, maintenance, pembelian, aset) selalu '
      'terhubung ke properti yang benar — jadi laporan pemilik selalu akurat dan konsisten.',
      ['Tidak perlu tindakan manual — cukup pastikan setiap properti sudah memiliki Analytic Account '
       '(dibuat sekali lewat tombol "Create Analytic Account" pada form properti).'])
doc.add_page_break()

fitur('Aset Tetap & Penyusutan', 'rental_management_asset',
      'Pencatatan aset tetap milik properti (mesin, peralatan, dsb.), penyusutan nilainya dari waktu '
      'ke waktu, dan penilaian ulang (revaluasi) bila diperlukan.',
      'Beban penyusutan tercatat otomatis dan rapi setiap bulan, serta ikut masuk sebagai beban dalam '
      'laporan keuangan properti terkait.',
      ['Daftarkan aset baru: nilai perolehan, metode penyusutan, dan jangka waktu penyusutan.',
       'Klik "Compute Depreciation" untuk membuat jadwal penyusutan otomatis.',
       'Biarkan sistem memposting beban penyusutan tiap periode secara otomatis (atau lakukan manual '
       'lewat tombol Post Due Depreciation).',
       'Bila nilai aset perlu disesuaikan (naik/turun), gunakan menu Revaluation.'])

fitur('Pajak & e-Faktur (CORETAX)', 'rental_management_coretax',
      'Menyiapkan data pajak (faktur pajak, retur, dan lampiran SPT) dalam format yang bisa diunggah '
      'ke portal CORETAX milik DJP.',
      'Kewajiban pelaporan pajak lebih tertata karena data faktur & lampiran SPT bisa ditarik langsung '
      'dari transaksi yang sudah tercatat di sistem, mengurangi kesalahan input ulang.',
      ['Lengkapi data NPWP/NIK dan ID TKU pada profil perusahaan dan tiap pelanggan.',
       'Pastikan kode pajak pada produk (sewa, jasa) sudah sesuai.',
       'Buka menu CORETAX → Export, pilih jenis dokumen (Faktur Keluaran, Retur, SPT, dst.) dan periode.',
       'Unduh file yang dihasilkan untuk diproses lebih lanjut pada portal CORETAX resmi.'])
doc.add_page_break()

# ==================== BAB 9 — LAYANAN MANDIRI ====================
bab(9, 'Layanan Mandiri untuk Tenant & Pemilik')
fitur('Portal Tenant', 'rental_management_portal',
      'Halaman khusus bagi tenant untuk melihat kontraknya sendiri dan tagihan yang harus dibayar, '
      'tanpa perlu menghubungi tim operasional untuk hal-hal dasar.',
      'Mengurangi beban tim operasional untuk pertanyaan rutin, dan tenant merasa lebih transparan '
      'karena bisa memantau sendiri statusnya.',
      ['Berikan akses portal ke kontak tenant (Contacts → pilih kontak → Grant Portal Access).',
       'Tenant login dan membuka menu "Contracts" untuk melihat daftar kontrak & detailnya.',
       'Dari detail kontrak, tenant bisa langsung melihat daftar invoicenya.'])

fitur('Portal Pemilik (Owner Portal)', 'rental_management_owner_portal',
      'Halaman khusus bagi pemilik properti untuk melihat properti miliknya dan riwayat penyaluran '
      'dana (remittance) yang sudah diterima.',
      'Pemilik mendapat transparansi penuh atas propertinya tanpa harus menunggu laporan dikirim manual.',
      ['Berikan akses portal ke kontak pemilik.',
       'Pemilik login dan membuka menu "Properties" untuk melihat daftar properti miliknya.',
       'Dari sana, pemilik dapat melihat riwayat remittance yang sudah diterima.'])

fitur('Dokumen Kontrak & Tanda Tangan Elektronik', 'rental_management_documents / _document_ce + _esign',
      'Penyimpanan dokumen terkait properti (kontrak, sertifikat, izin) secara terorganisir, serta '
      'pelacakan proses tanda tangan elektronik kontrak.',
      'Dokumen penting tidak lagi tersebar di email/folder pribadi; semua tersimpan rapi per properti, '
      'dan status tanda tangan kontrak mudah dipantau.',
      ['Lampirkan dokumen langsung pada form properti/kontrak terkait.',
       'Untuk penandatanganan, klik "Request Signature" pada kontrak, lalu tandai "Mark Signed" setelah '
       'dokumen ditandatangani.'])
doc.add_page_break()

# ==================== BAB 10 — RISIKO ====================
bab(10, 'Mengelola Risiko')
fitur('Asuransi Gedung', 'rental_management_insurance',
      'Pencatatan polis asuransi yang melindungi gedung/properti, lengkap dengan pengingat otomatis '
      'menjelang tanggal kedaluwarsa polis.',
      'Properti selalu terlindungi asuransi tanpa risiko polis kedaluwarsa tanpa disadari.',
      ['Catat polis: penanggung, nomor polis, nilai pertanggungan, dan periode berlaku.',
       'Sistem otomatis mengingatkan tim terkait menjelang tanggal kedaluwarsa untuk perpanjangan.'])

fitur('Valuasi Properti', 'rental_management_valuation',
      'Pencatatan hasil penilaian (appraisal) nilai pasar properti secara berkala.',
      'Manajemen memiliki riwayat nilai pasar properti dari waktu ke waktu untuk keperluan keputusan '
      'bisnis maupun pelaporan ke pemilik.',
      ['Catat hasil valuasi baru: tanggal, metode penilaian, penilai, dan nilai pasar.',
       'Nilai valuasi terbaru otomatis menjadi acuan nilai terkini pada data properti.'])
doc.add_page_break()

# ==================== BAB 11 — AKHIR SIKLUS SEWA ====================
bab(11, 'Perpanjangan Sewa & Akhir Masa Kontrak')
fitur('Pengingat Berakhirnya Sewa & Perpanjangan', 'rental_management_lease_expiry',
      'Daftar kontrak yang akan berakhir dalam 30/90 hari mendatang, lengkap dengan pengingat otomatis '
      'agar tim segera menindaklanjuti perpanjangan.',
      'Tidak ada kontrak yang "kelewat" tanpa tindak lanjut, sehingga risiko unit kosong mendadak berkurang.',
      ['Buka menu Lease Expiry, gunakan filter 30 hari atau 90 hari untuk melihat kontrak yang mendekati akhir.',
       'Tindak lanjuti sesuai keputusan: perpanjang (aktifkan rent escalation bila relevan) atau siapkan proses move-out.'])

fitur('Papan Ketersediaan Unit (Vacancy)', 'rental_management_vacancy',
      'Tampilan ringkas seluruh unit yang dikelompokkan berdasarkan statusnya: tersedia, sedang '
      'disewa, atau sudah dibooking.',
      'Tim leasing dapat langsung melihat unit mana yang bisa segera dipasarkan tanpa perlu memeriksa '
      'satu per satu.',
      ['Buka menu Vacancy / Availability Board.',
       'Gunakan filter status ("Available") untuk melihat unit yang siap dipasarkan ke calon tenant baru.'])

para('Untuk proses serah-terima saat tenant keluar (Move-out) dan penyelesaian deposit, lihat kembali '
     'Bab 5 (Handover) dan Bab 4 (Security Deposit) — keduanya dipakai bersamaan saat kontrak berakhir.',
     italic=True, color=GREY)
doc.add_page_break()

# ==================== BAB 12 — DASHBOARD ====================
bab(12, 'Memantau Semuanya lewat Dashboard')
fitur('Dashboard KPI Manajemen', 'rental_management_dashboard',
      'Ringkasan angka-angka penting dalam satu tampilan: jumlah properti, kontrak aktif, pendapatan '
      'bersih (NOI), tingkat tunggakan, tingkat penagihan (collection rate), dan kontrak yang akan '
      'berakhir dalam 12 bulan ke depan.',
      'Manajemen tidak perlu membuka banyak laporan terpisah untuk mengetahui kondisi bisnis secara '
      'umum — semua indikator penting tersaji dalam satu halaman.',
      ['Buka menu KPI Dashboard.',
       'Pilih periode dan/atau properti yang ingin dipantau.',
       'Tinjau kartu-kartu KPI; gunakan sebagai bahan diskusi rutin dengan tim.'])

# ==================== BAB 13 — RINGKASAN & REFERENSI ====================
doc.add_page_break()
bab(13, 'Ringkasan Peta Fitur (Referensi Cepat)')
para('Tabel berikut merangkum seluruh fitur pada panduan ini beserta menu tempat mengaksesnya, '
     'untuk memudahkan pencarian cepat.')
ref = doc.add_table(rows=1, cols=3); ref.style = 'Table Grid'
for i, htxt in enumerate(['Bab & Fitur', 'Menu Akses', 'Ringkasan Manfaat']):
    shade(ref.rows[0].cells[i], '1F3964'); ct(ref.rows[0].cells[i], htxt, bold=True, white=True, size=9)
REF_ROWS = [
    ('2. Properti & Unit', 'Properties → Properties', 'Pondasi data seluruh transaksi'),
    ('2. Kepemilikan', 'form Properti → Owners & Financial', 'Dasar pembagian dana ke pemilik'),
    ('3. Pipeline Leasing', 'CRM', 'Kelola prospek → kontrak otomatis'),
    ('4. Kontrak Sewa', 'form Tenancy', 'Sumber kebenaran penagihan'),
    ('4. Casual Leasing', 'Casual Leases', 'Sewa jangka pendek'),
    ('4. Guarantee', 'Tenant Guarantees', 'Jaminan + alert kedaluwarsa'),
    ('4. Security Deposit', 'Financial Reports → Security Deposits', 'Deposit sebagai liabilitas'),
    ('4. Rent Escalation', 'tab kontrak', 'Kenaikan sewa otomatis'),
    ('5. Handover', 'Handovers', 'Checklist serah-terima unit'),
    ('5. Access & Parking', 'Access & Visitors / Parking', 'Kontrol akses & petak parkir'),
    ('6. GTO & Meter', 'Leasing Operations', 'Sewa omzet & recharge utilitas'),
    ('6. CAM', 'CAM / Service Charge', 'Biaya bersama proporsional'),
    ('6. Dunning', 'Dunning Levels', 'Penagihan otomatis bertingkat'),
    ('7. Maintenance & PPM', 'Maintenance', 'Perbaikan & perawatan terjadwal'),
    ('7. Purchase', 'Property Purchase Orders', 'Pengadaan terhubung properti'),
    ('8. Report Category', 'Chart of Accounts', 'Pemetaan akun ke laporan'),
    ('8. Budget', 'Property Budgets', 'Anggaran vs realisasi'),
    ('8. Trust & Remittance', 'Owner Remittances', 'Penyaluran dana ke pemilik'),
    ('8. Owners Statement', 'Owners Statement', 'Laporan lengkap 9 bagian'),
    ('8. Fixed Asset', 'Fixed Assets', 'Penyusutan & revaluasi'),
    ('8. CORETAX', 'CORETAX', 'Ekspor faktur pajak & SPT'),
    ('9. Tenant Portal', '/my (tenant)', 'Self-service tenant'),
    ('9. Owner Portal', '/my/properties (owner)', 'Self-service pemilik'),
    ('9. Dokumen & e-Sign', 'form Properti/Kontrak', 'Simpan dokumen & tanda tangan'),
    ('10. Insurance', 'Insurance Policies', 'Polis asuransi gedung'),
    ('10. Valuation', 'Property Valuations', 'Riwayat nilai pasar'),
    ('11. Lease Expiry', 'Lease Expiry', 'Reminder perpanjangan sewa'),
    ('11. Vacancy', 'Vacancy / Availability', 'Papan status unit'),
    ('12. Dashboard', 'KPI Dashboard', 'Ringkasan KPI manajemen'),
]
for a, b_, c in REF_ROWS:
    cells = ref.add_row().cells
    ct(cells[0], a, bold=True, size=8.5); ct(cells[1], b_, size=8.5); ct(cells[2], c, size=8.5)
    cells[0].width = Inches(1.9); cells[1].width = Inches(2.1); cells[2].width = Inches(2.7)
doc.add_paragraph()

para(' ')
para('Catatan: Panduan ini bersifat fungsional untuk pengguna sehari-hari. Untuk detail teknis '
     '(nama model, field, hook integrasi) lihat dokumen Blueprint Instalasi & Konfigurasi. Untuk '
     'skenario pengujian, lihat UAT Tracker.', italic=True, color=GREY, size=9.5)

doc.save('../../docs/Panduan_Pengguna_Fitur_Sesuai_Alur_Bisnis.docx')
print('SAVED OK')
