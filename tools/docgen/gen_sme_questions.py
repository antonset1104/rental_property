# -*- coding: utf-8 -*-
"""Generator kuesioner SME komprehensif — mencakup seluruh 26 modul PMS Odoo 19."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x39, 0x64)
GREEN = RGBColor(0x00, 0x6A, 0x4E)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

# --- Landscape ---
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = Inches(0.6); sec.right_margin = Inches(0.6)
sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6)
USABLE = 10.1  # inch


def shade(c, h):
    tcPr = c._tc.get_or_add_tcPr(); s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), h)
    tcPr.append(s)


def ct(c, t, bold=False, white=False, size=9, color=None):
    c.text = ''; p = c.paragraphs[0]; r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color: r.font.color.rgb = color


def h1(t):
    p = doc.add_heading(level=1); r = p.add_run(t); r.font.color.rgb = NAVY; r.font.size = Pt(15); return p


def h2(t):
    p = doc.add_heading(level=2); r = p.add_run(t); r.font.color.rgb = GREEN; r.font.size = Pt(12.5); return p


def para(t, bold=False, italic=False, size=10.5, color=None, after=6):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p


def bullet(t, pre=None):
    p = doc.add_paragraph(style='List Bullet')
    if pre:
        rr = p.add_run(pre); rr.bold = True
    p.add_run(t); return p


# kolom: No | Pertanyaan | Tujuan & Konteks | Jawaban SME
COLW = [Inches(0.4), Inches(3.7), Inches(2.9), Inches(3.1)]
COUNTER = {'n': 0}


def qsection(letter, title, rows):
    h2('%s. %s' % (letter, title))
    t = doc.add_table(rows=1, cols=4); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    heads = ['No', 'Pertanyaan untuk Key User / SME', 'Tujuan & Konteks', 'Jawaban SME']
    for i, htx in enumerate(heads):
        shade(t.rows[0].cells[i], '1F3964'); ct(t.rows[0].cells[i], htx, bold=True, white=True, size=9)
        t.rows[0].cells[i].width = COLW[i]
    for q, ctx in rows:
        COUNTER['n'] += 1
        cells = t.add_row().cells
        ct(cells[0], str(COUNTER['n']), size=8.5)
        ct(cells[1], q, size=8.5)
        ct(cells[2], ctx, size=8, color=GREY)
        ct(cells[3], '', size=8.5)
        for i, w in enumerate(COLW):
            cells[i].width = w
    doc.add_paragraph()


# ==================== COVER ====================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('DAFTAR PERTANYAAN KEY USER / SME — KOMPREHENSIF')
r.bold = True; r.font.size = Pt(19); r.font.color.rgb = NAVY
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Requirement Gathering seluruh modul Property Management System Odoo 19\n'
               '(addon rental_management + 26 modul companion) — untuk Blueprint lengkap')
r.font.size = Pt(12.5); r.bold = True; r.font.color.rgb = GREEN
doc.add_paragraph()
meta = doc.add_table(rows=6, cols=2); meta.style = 'Table Grid'
for i, (k, v) in enumerate([
    ('Dokumen', 'Functional & Technical Questionnaire (seluruh modul)'),
    ('Platform', 'Odoo 19 Community + addon rental_management v3.3.9 (TechKhedut)'),
    ('Peserta diskusi', 'IT / System Analyst ↔ User Operasional & Accounting'),
    ('Tujuan', 'Mengumpulkan aturan bisnis untuk menyusun Blueprint & konfigurasi sistem'),
    ('Cara pengisian', 'Isi kolom "Jawaban SME" sespesifik mungkin; sertakan contoh nilai/aturan'),
    ('Tanggal', '6 Juli 2026'),
]):
    shade(meta.rows[i].cells[0], 'E7ECF5'); ct(meta.rows[i].cells[0], k, bold=True, size=10)
    ct(meta.rows[i].cells[1], v, size=10)
    meta.rows[i].cells[0].width = Inches(2.2); meta.rows[i].cells[1].width = Inches(7.6)
doc.add_paragraph()

h1('Petunjuk Pengisian')
bullet('Setiap baris memiliki kolom "Tujuan & Konteks" agar SME memahami mengapa data diperlukan.', pre='Konteks: ')
bullet('Isi kolom "Jawaban SME" langsung. Bila belum pasti, tulis asumsi sementara + PIC penanggung jawab.', pre='Jawaban: ')
bullet('Jawaban akan menjadi dasar Functional Spec, konfigurasi sistem, dan Blueprint final.', pre='Output: ')
bullet('Pertanyaan dikelompokkan per area fungsional/modul (A–AF). Lewati bagian yang tidak relevan dengan bisnis Anda dan tandai "N/A".', pre='Lingkup: ')
bullet('Bagian bertanda ★ paling kritikal untuk pelaporan keuangan (Owners Statement).', pre='Prioritas: ')
bullet('Dokumen ini TIDAK mengulang ~40 pertanyaan pada "Daftar Pertanyaan Key User – Laporan Keuangan" '
       'sebelumnya (area A–Q seputar Owners Statement) yang sudah dijawab — mis. frekuensi laporan bulanan, '
       'mata uang IDR, tidak ada rekening trust terpisah, remittance proporsional kepemilikan, struktur kode '
       'akun, budget disusun Property Manager dgn approval 3 level, PPN 11%/PPh 10%, deposit sebagai titipan. '
       'Beberapa pertanyaan di bawah adalah tindak lanjut (follow-up) dari jawaban tersebut.', pre='Catatan: ')
doc.add_page_break()

h1('Bagian I — Umum, Master Data & Kontrak')

# ==================== A. Profil & Ruang Lingkup ====================
qsection('A', 'Profil Bisnis & Ruang Lingkup', [
    ('Jenis properti yang dikelola (mal/ritel, perkantoran, apartemen, pergudangan, mixed-use)? Sebutkan komposisinya.',
     'Menentukan model unit, tipe sewa, dan fitur yang relevan.'),
    ('Apakah pengelolaan atas nama pemilik pihak ketiga (managed property) atau milik sendiri, atau campuran?',
     'Menentukan perlunya Trust Accounting & Owners Statement.'),
    ('Berapa entitas legal/PT yang terlibat? Apakah perlu multi-company di Odoo?',
     'Menentukan struktur multi-company & konsolidasi.'),
    ('Siapa saja peran pengguna sistem (operasional, accounting, manajemen, pemilik, tenant) dan tanggung jawabnya?',
     'Menentukan hak akses (security groups) & portal.'),
    ('Sistem apa yang dipakai saat ini (IFCA/Excel/lainnya)? Fitur mana yang wajib disamai/ditingkatkan?',
     'Menentukan cakupan migrasi & gap fungsional.'),
])

# ==================== B. Master Properti & Kepemilikan ====================
qsection('B', 'Master Data Properti, Unit & Kepemilikan ★', [
    ('Bagaimana hierarki lokasi: Proyek/Gedung → Lantai → Unit? Ada penomoran/standar kode unit?',
     'Menentukan struktur property.details & pengelompokan laporan.'),
    ('Atribut unit yang wajib dicatat (luas sewa/semi gross/net, NLA, zona, peruntukan)?',
     'Menentukan field unit & dasar apportionment CAM.'),
    ('Data pemilik yang perlu dicatat (nama, NPWP, rekening bank, kontak, porsi, tanggal efektif)?',
     'Menentukan master owner & data remittance.'),
    ('Untuk perubahan kepemilikan di tengah periode (sudah dikonfirmasi "bisa terjadi"): bagaimana pembagian '
     'hasil di bulan transisi dihitung — pro-rata harian, atau efektif awal bulan berikutnya?',
     'Menentukan aturan pro-rata teknis pada remittance & statement.'),
])

# ==================== C. Kontrak & Leasing ====================
qsection('C', 'Kontrak Sewa & Leasing (rental_management)', [
    ('Jenis kontrak yang ada (fixed term, evergreen, short-term/casual, anchor vs specialty)?',
     'Menentukan tipe tenancy & aturan tenor.'),
    ('Komponen tagihan dalam satu kontrak (base rent, service charge/CAM, GTO, utilitas, parkir, promosi)?',
     'Menentukan struktur invoice & produk.'),
    ('Frekuensi & timing penagihan (bulanan di muka/belakang, kuartalan)? Tanggal jatuh tempo standar?',
     'Menentukan penjadwalan invoice & aging.'),
    ('Aturan pro-rata untuk move-in/move-out di tengah bulan?',
     'Menentukan perhitungan sewa parsial.'),
    ('Apakah ada rent-free period, fit-out period, atau stepped rent?',
     'Menentukan penjadwalan & pengakuan pendapatan.'),
    ('Ketentuan pajak atas sewa (PPN 11%, PPh 4(2) final 10%)? Siapa yang memungut/menyetor?',
     'Menentukan tax setup & CORETAX.'),
    ('Dokumen apa yang melekat pada kontrak (PKS, KTP/NPWP, jaminan)?',
     'Menentukan DMS & e-sign.'),
])

# ==================== D. CRM / Pipeline Leasing ====================
qsection('D', 'Pipeline Leasing / CRM (rental_management_crm)', [
    ('Tahapan pipeline leasing dari prospek hingga deal (mis. Inquiry → Viewing → Nego → LOI → Signed)?',
     'Menentukan stages CRM & konversi ke kontrak.'),
    ('Sumber prospek (walk-in, broker, online, referral) yang perlu dilacak?',
     'Menentukan lead source & pelaporan konversi.'),
    ('Data minimum saat membuat prospek (unit diminati, target rent, durasi, target move-in)?',
     'Menentukan field lead & prefilling kontrak.'),
    ('Apakah ada target/KPI leasing (occupancy, waktu deal, WALE) yang ingin dipantau?',
     'Menentukan dashboard & metrik.'),
])

doc.add_page_break()
h1('Bagian II — Keuangan, Laporan & Pajak ★')

# ==================== E. Owners Statement ====================
qsection('E', 'Owners Statement & Format Laporan ★', [
    ('Sembilan section (Performance Summary, I&E, Receipts & Payments, Tenant Balances, Aged Arrears, Payment Details, Trial Balance, Balance Sheet, GST) — semuanya diperlukan? Ada tambahan?',
     'Menentukan cakupan report suite.'),
    ('Perlu narasi/ringkasan eksekutif (catatan manajemen) selain angka pada tiap statement?',
     'Menentukan kebutuhan free-text section pada laporan.'),
    ('Untuk statement konsolidasi lintas properti milik satu owner (sudah dikonfirmasi wajib): apakah tetap '
     'menampilkan rincian per properti, atau hanya total gabungan?',
     'Menentukan struktur laporan konsolidasi.'),
])

# ==================== F. CoA & Report Category ====================
qsection('F', 'Chart of Accounts & Kategori Laporan ★', [
    ('Perlu sub-group/hierarki dalam section (mis. Utilities → Listrik, Air)?',
     'Menentukan struktur parent kategori.'),
    ('Kapan tabel mapping resmi akun→kategori (Statutory/Variable/Direct Recharge/Non-Recoverable/Capex, '
     'sudah diminta pada sesi sebelumnya) dapat disediakan tim Accounting? Siapa PIC-nya?',
     'Menentukan jadwal & PIC penyediaan mapping sebelum konfigurasi sistem.'),
])

# ==================== G. Trust & Remittance ====================
qsection('G', 'Trust Accounting & Owner Remittance ★', [
    ('Karena dikonfirmasi TIDAK ada rekening trust/titipan terpisah (dana diterima langsung ke rekening '
     'operasional): apakah cukup memakai satu akun Bank/Trust "virtual" (GL/analytic) untuk memantau saldo '
     'per properti, tanpa rekening bank fisik terpisah?',
     'Menentukan apakah modul Trust Accounting perlu disederhanakan menjadi saldo GL saja (bukan rekonsiliasi bank).'),
    ('Dana yang harus diteruskan ke pemilik ditransfer dari rekening operasional mana, dan siapa yang '
     'mengeksekusi transfer (finance/treasury)?',
     'Menentukan proses & approval eksekusi remittance.'),
])

# ==================== H. Performance Summary & Budget ====================
qsection('H', 'Performance Summary, Budget & Variance ★', [
    ('Perlu tampilan Actual vs Budget vs Variance vs %Var? Untuk Period & YTD?',
     'Menentukan kolom & model budget.'),
    ('Basis Performance Summary: accrual, cash, atau keduanya (seperti PDF)?',
     'Menentukan sumber data & rekonsiliasi.'),
    ('Untuk approval budget 3 level (sudah dikonfirmasi): siapa role di tiap level (mis. Property Manager → '
     'Finance Manager → Direksi) dan berapa ambang nilai (jika ada) per level?',
     'Menentukan matriks approval budget di sistem.'),
])

# ==================== I. Tenant Balances & Arrears ====================
qsection('I', 'Tenant Balances & Aged Arrears', [
    ('Bagaimana urutan alokasi pembayaran ke saldo tenant (oldest-first/FIFO, per invoice, atau per kategori charge)?',
     'Menentukan logika reconciliation & saldo berjalan.'),
    ('Apakah bunga/denda keterlambatan diterapkan pada tunggakan? Berapa % / nominal dan sejak hari ke berapa?',
     'Menentukan konfigurasi dunning & late fee.'),
    ('Bagaimana perlakuan tenant in-credit (saldo lebih bayar) — dikompensasi ke tagihan berikutnya atau direfund? (pertanyaan ini belum terjawab pada sesi sebelumnya)',
     'Menentukan logika "Net Arrears" & baris in-credit.'),
])

# ==================== J. Payment Details & AP ====================
qsection('J', 'Payment Details / Pembayaran Vendor', [
    ('Detail pembayaran vendor perlu dikelompokkan per kode akun/kategori di statement?',
     'Menentukan Payment Details report.'),
    ('Perlu lampiran bukti bayar/tanggal kliring pada Payment Details?',
     'Menentukan detail & attachment.'),
])

# ==================== K. GST/PPN & Pajak ====================
qsection('K', 'GST/PPN & Rekonsiliasi Pajak ★', [
    ('Apakah ada withholding PPh 23 yang dipungut tenant (bukti potong) di luar PPh final 10% yang perlu dicatat?',
     'Menentukan pencatatan WHT tambahan.'),
    ('Bukti potong pajak dari tenant/vendor — perlu diunggah & disimpan di sistem, atau cukup dicatat nomornya?',
     'Menentukan kebutuhan lampiran dokumen pajak.'),
])

doc.add_page_break()
h1('Bagian III — Modul Operasional & Charge')

# ==================== L. GTO ====================
qsection('L', 'GTO / Revenue Sharing (rental_management_gto_meter)', [
    ('Apakah ada sewa berbasis omzet (GTO)? Model mana: Higher-of, Base+Overage, atau Pure %?',
     'Menentukan konfigurasi GTO di kontrak.'),
    ('Persentase GTO & breakpoint per kategori tenant? Contoh angka?',
     'Menentukan tarif & breakpoint.'),
    ('Bagaimana omzet dideklarasikan (self-report tenant, POS integration, audit)? Frekuensi?',
     'Menentukan proses turnover declaration.'),
    ('Apakah omzet dikoreksi/rekonsiliasi tahunan (annual reconciliation)?',
     'Menentukan penyesuaian akhir tahun.'),
])

# ==================== M. Meter & Utility ====================
qsection('M', 'Meter & Utility Recharge (rental_management_gto_meter)', [
    ('Utilitas apa yang di-recharge ke tenant (listrik, air, gas, chilled water)? Ada sub-meter?',
     'Menentukan master meter & produk.'),
    ('Tarif per unit (flat, tiered, mengikuti PLN/PDAM + margin)? Ada minimum charge?',
     'Menentukan perhitungan recharge.'),
    ('Frekuensi pembacaan meter & siapa yang input?',
     'Menentukan proses reading & jadwal.'),
    ('Apakah ada standing charge/abonemen terpisah dari pemakaian?',
     'Menentukan komponen tagihan.'),
])

# ==================== N. CAM ====================
qsection('N', 'CAM / Service Charge (rental_management_cam)', [
    ('Dasar apportionment service charge (pro-rata luas/NLA, rate per m², fixed)?',
     'Menentukan metode alokasi CAM.'),
    ('Komponen biaya pool CAM (kebersihan, keamanan, listrik area umum, dll)?',
     'Menentukan expense pool.'),
    ('Apakah service charge ditagih atas dasar budget (estimasi) lalu direkonsiliasi ke actual?',
     'Menentukan siklus budget vs actual & balancing charge.'),
    ('Frekuensi rekonsiliasi CAM (tahunan)? Bagaimana selisih dibebankan/dikreditkan?',
     'Menentukan reconciliation.'),
])

# ==================== O. Rent Escalation ====================
qsection('O', 'Rent Escalation (rental_management_rent_escalation)', [
    ('Aturan kenaikan sewa (fixed % per tahun, nominal, mengikuti indeks/CPI)?',
     'Menentukan tipe & nilai escalation.'),
    ('Kapan kenaikan berlaku (anniversary kontrak, awal tahun kalender)?',
     'Menentukan tanggal efektif & cron.'),
    ('Apakah kenaikan perlu persetujuan/pemberitahuan ke tenant dulu?',
     'Menentukan approval/notifikasi.'),
])

# ==================== P. Casual Leasing ====================
qsection('P', 'Casual Leasing / Short-term (rental_management_casual_leasing)', [
    ('Apakah ada sewa jangka pendek (booth, pop-up, atrium, media)? Basis tarif (harian/mingguan/fixed)?',
     'Menentukan casual lease & perhitungan.'),
    ('Apakah butuh deposit/izin khusus untuk casual lease?',
     'Menentukan syarat & dokumen.'),
    ('Bagaimana pendapatan casual dilaporkan (masuk properti terkait)?',
     'Menentukan atribusi & laporan.'),
])

# ==================== Q. Parking ====================
qsection('Q', 'Parking Management (rental_management_parking)', [
    ('Apakah parkir dikelola sistem ini? Jenis (bulanan tenant, harian umum, reserved)?',
     'Menentukan model bay & tarif.'),
    ('Tarif parkir & apakah ditagih via kontrak tenant atau terpisah?',
     'Menentukan invoicing parkir.'),
    ('Perlu alokasi bay spesifik ke tenant/nomor kendaraan?',
     'Menentukan alokasi & kartu akses.'),
])

# ==================== R. Handover / Fit-out ====================
qsection('R', 'Handover: Move-in / Fit-out / Move-out (rental_management_handover + project)', [
    ('Checklist standar untuk Move-in, Fit-out, dan Move-out? Item wajibnya?',
     'Menentukan template checklist.'),
    ('Untuk Fit-out: perlu kontraktor, fit-out bond, izin, dan periode fit-out?',
     'Menentukan field & bond.'),
    ('Apakah fit-out dikelola sebagai proyek (task, timeline)?',
     'Menentukan integrasi Project.'),
    ('Untuk Move-out: aturan make-good/reinstatement & pengembalian kunci/akses?',
     'Menentukan proses & penyelesaian deposit.'),
])

doc.add_page_break()
h1('Bagian IV — Risiko, Maintenance, Aset & Procurement')

# ==================== S. Security Deposit ====================
qsection('S', 'Security Deposit (rental_management_financial_report)', [
    ('Berapa besaran deposit standar (mis. 3 bulan sewa) & komponen apa saja yang dijamin?',
     'Menentukan nilai & cakupan deposit.'),
    ('Apakah deposit berupa uang tunai, bank guarantee, atau keduanya?',
     'Menentukan jenis jaminan (link ke Guarantee).'),
    ('Apakah deposit dikreditkan bunga atau di-top-up saat sewa naik (mis. akibat rent escalation)?',
     'Menentukan perlakuan deposit terhadap eskalasi.'),
    ('Berapa lama proses settlement & BAST (sudah dikonfirmasi sebagai syarat refund) biasanya berlangsung setelah sewa berakhir?',
     'Menentukan SLA penyelesaian move-out & refund deposit.'),
])

# ==================== T. Guarantee ====================
qsection('T', 'Bank / Insurance Guarantee (rental_management_guarantee)', [
    ('Jenis jaminan yang diterima (bank guarantee, insurance bond, security guarantee)?',
     'Menentukan tipe guarantee.'),
    ('Data yang dilacak (penerbit, nomor, nilai, terbit, expiry) & lead reminder sebelum kedaluwarsa?',
     'Menentukan field & alert.'),
    ('Proses saat guarantee expiring/klaim/release — siapa PIC?',
     'Menentukan workflow & tanggung jawab.'),
])

# ==================== U. Insurance ====================
qsection('U', 'Insurance Polis Gedung (rental_management_insurance)', [
    ('Polis asuransi apa yang dikelola (property all risk, public liability, business interruption)?',
     'Menentukan register polis.'),
    ('Data polis (penanggung, nomor, nilai pertanggungan, premi, periode) & reminder perpanjangan?',
     'Menentukan field & alert expiry.'),
    ('Apakah premi asuransi termasuk komponen recoverable ke tenant (via CAM)?',
     'Menentukan recovery & mapping akun.'),
])

# ==================== V. Dunning ====================
qsection('V', 'Dunning / Penagihan (rental_management_dunning)', [
    ('Tingkatan penagihan (mis. H+7 reminder, H+14 SP1, H+30 SP2) & isi/template komunikasinya?',
     'Menentukan dunning ladder & template.'),
    ('Apakah ada denda keterlambatan (late fee)? Basis (% per bulan, flat)?',
     'Menentukan produk late fee & perhitungan.'),
    ('Kapan tindakan lanjutan (pemutusan, pencairan jaminan) diambil?',
     'Menentukan eskalasi.'),
    ('Siapa penerima notifikasi (tenant, internal collection)?',
     'Menentukan email & PIC.'),
])

# ==================== W. Lease Expiry & Vacancy ====================
qsection('W', 'Lease Expiry, Renewal & Vacancy (lease_expiry + vacancy)', [
    ('Kapan proses perpanjangan dimulai sebelum jatuh tempo (mis. 90/60/30 hari)?',
     'Menentukan reminder renewal.'),
    ('Siapa yang ditugaskan menindaklanjuti expiry & bagaimana keputusan renew/keluar dicatat?',
     'Menentukan aktivitas & PIC.'),
    ('Definisi status unit (available, booked, on-lease) untuk papan vacancy?',
     'Menentukan pemetaan status.'),
    ('Metrik okupansi/vacancy yang ingin dipantau (rate, WALE)?',
     'Menentukan dashboard.'),
])

# ==================== X. Maintenance & PPM ====================
qsection('X', 'Maintenance & Preventive Maintenance (maintenance + rental_management_ppm)', [
    ('Kategori maintenance (corrective/reaktif vs preventive) & aset/peralatan yang dijadwalkan PPM?',
     'Menentukan master aset & plan PPM.'),
    ('Frekuensi PPM per jenis peralatan (lift, genset, AHU, fire system) & SLA penyelesaian?',
     'Menentukan jadwal & SLA.'),
    ('Apakah biaya maintenance bisa di-recharge ke tenant? Kondisinya?',
     'Menentukan recharge & recovery.'),
    ('Alur permintaan → penugasan → penyelesaian → verifikasi; siapa PIC tiap tahap?',
     'Menentukan workflow maintenance.'),
])

# ==================== Y. Procurement ====================
qsection('Y', 'Procurement / Purchase (rental_management_purchase)', [
    ('Untuk alur PR/PO/GRS yang sudah dikonfirmasi terintegrasi: berapa ambang nilai per level approval, dan siapa approver di tiap level?',
     'Menentukan matriks approval purchase.'),
    ('Apakah PO dikaitkan ke properti/kontrak/maintenance untuk atribusi biaya?',
     'Menentukan link & analytic.'),
    ('Manajemen vendor (master, term pembayaran, evaluasi)?',
     'Menentukan master vendor.'),
])

# ==================== Z. Fixed Asset ====================
qsection('Z', 'Fixed Asset, Depreciation & Revaluation (rental_management_asset)', [
    ('Kelas aset & metode penyusutan (garis lurus/saldo menurun), umur manfaat per kelas?',
     'Menentukan konfigurasi board penyusutan.'),
    ('Kebijakan revaluasi aset (frekuensi, dasar, akun cadangan revaluasi)?',
     'Menentukan proses revaluation.'),
    ('Apakah penyusutan dialokasikan per properti (analytic) & masuk Owners Statement?',
     'Menentukan atribusi beban penyusutan.'),
    ('Apakah register aset perlu sinkron ke lampiran pajak (CORETAX L9)?',
     'Menentukan integrasi L9.'),
])

# ==================== AA. Valuation ====================
qsection('AA', 'Property Valuation (rental_management_valuation)', [
    ('Frekuensi valuasi pasar (tahunan) & metode (income/market/cost)?',
     'Menentukan register valuasi.'),
    ('Siapa penilai (internal/KJPP) & data yang dicatat (nilai, tanggal, cap rate)?',
     'Menentukan field valuasi.'),
])

doc.add_page_break()
h1('Bagian V — Pajak Digital, Portal, Dokumen & Teknis')

# ==================== AB. CORETAX ====================
qsection('AB', 'CORETAX / e-Faktur & SPT (rental_management_coretax) ★', [
    ('Sudah dikonfirmasi nomor faktur dari CORETAX "dimasukkan ke sistem" (input manual). Apakah artinya faktur '
     'tetap diterbitkan lewat portal CORETAX lalu nomornya dicatat balik ke invoice Odoo, ATAU justru Odoo yang '
     'akan meng-generate XML untuk diunggah ke CORETAX (seperti modul ini dibangun)? Mohon konfirmasi arah proses.',
     'Menentukan arah integrasi: import nomor faktur vs export XML e-Faktur — berdampak besar ke desain modul.'),
    ('Volume faktur pajak (keluaran & masukan) per bulan saat ini?',
     'Menentukan kebutuhan & skala ekspor/import.'),
    ('Data wajib pajak yang tersedia (NPWP 16 digit/NIK, ID TKU) untuk company & pelanggan?',
     'Menentukan kelengkapan master pajak.'),
    ('Transaction code faktur yang umum dipakai (04 DPP nilai lain, 01, dst) untuk sewa?',
     'Menentukan default TrxCode.'),
    ('Lampiran SPT mana yang relevan (L9, L3B, L11A Uncollectible/NonPerforming/Promosi/Entertainment, L10A)?',
     'Menentukan register SPT yang dipakai.'),
    ('Apakah ada retur faktur (kredit) yang perlu diekspor (Retur PM/Lampiran C)?',
     'Menentukan ekspor retur.'),
    ('Siapa yang bertanggung jawab ekspor & upload ke portal CORETAX + validasi XSD?',
     'Menentukan PIC & kontrol.'),
])

# ==================== AC. Portal ====================
qsection('AC', 'Tenant Portal & Owner Portal (portal + owner_portal)', [
    ('Apa yang boleh dilihat tenant di portal (kontrak, invoice, saldo, permintaan maintenance)?',
     'Menentukan cakupan & hak akses tenant.'),
    ('Apa yang boleh dilihat pemilik (properti, Owners Statement, remittance)?',
     'Menentukan cakupan Owner Portal.'),
    ('Apakah tenant boleh mengajukan permintaan/komplain lewat portal?',
     'Menentukan fitur self-service.'),
])

# ==================== AD. Documents & E-sign ====================
qsection('AD', 'Manajemen Dokumen & E-Signature (documents/document_ce + esign)', [
    ('Jenis dokumen yang disimpan per properti/kontrak & struktur folder/kategori?',
     'Menentukan DMS & kategori.'),
    ('Apakah pakai Odoo Enterprise Documents atau cukup lampiran Community?',
     'Menentukan modul dokumen yang dipakai.'),
    ('Apakah penandatanganan kontrak elektronik diperlukan? Lewat Odoo atau pihak ketiga (mis. Privy)?',
     'Menentukan e-sign & integrasi.'),
    ('Kebijakan retensi & akses dokumen (siapa boleh lihat/unduh)?',
     'Menentukan retensi & akses.'),
])

# ==================== AE. Access & Visitor ====================
qsection('AE', 'Access Card & Visitor Log (rental_management_access)', [
    ('Apakah kartu akses/parkir dikelola di sistem? Data yang dicatat (nomor, pemegang, masa berlaku)?',
     'Menentukan register kartu.'),
    ('Apakah buku tamu/visitor log diperlukan (check-in/out)? Terintegrasi hardware?',
     'Menentukan visitor log & integrasi.'),
])

# ==================== AF. Dashboard, Migrasi & Non-fungsional ====================
qsection('AF', 'KPI Dashboard, Migrasi Data & Non-Fungsional', [
    ('KPI utama untuk manajemen (NOI, collection rate, arrears, occupancy, WALE, expiring 12 bulan)?',
     'Menentukan isi dashboard.'),
    ('Untuk migrasi data historis (sudah dikonfirmasi "ada"): mohon detail per item — saldo awal tenant, deposit, '
     'budget berjalan, DAN register fixed asset & akumulasi penyusutan (belum disebutkan sebelumnya) — dalam format apa (Excel/ekspor MRI)?',
     'Melengkapi cakupan migrasi khususnya fixed asset yang belum dikonfirmasi, & menentukan format import.'),
    ('Tanggal cut-over/go-live yang ditargetkan & periode paralel run?',
     'Menentukan rencana go-live.'),
    ('Untuk periode yang sudah dikonfirmasi "dikunci" setelah laporan terbit: siapa saja yang punya hak override/backdate, dan lewat proses apa (approval khusus)?',
     'Menentukan hak akses & audit trail atas pengecualian kunci periode.'),
    ('Kebutuhan approval workflow lain (invoice, remittance, PO, budget) & batas nilainya?',
     'Menentukan matriks otorisasi.'),
    ('Ekspektasi laporan tambahan/kustom di luar Owners Statement (mis. rent roll, arrears summary)?',
     'Menentukan laporan tambahan.'),
    ('Kebutuhan integrasi eksternal (bank/host-to-host, virtual account, payment gateway, POS)?',
     'Menentukan integrasi & effort.'),
    ('SLA sistem, jam operasional, kebijakan backup & retensi data?',
     'Menentukan non-fungsional & infrastruktur.'),
])

# ==================== PENUTUP ====================
doc.add_page_break()
h1('Langkah Selanjutnya')
para('Setelah kolom "Jawaban SME" terisi, jawaban akan dikonsolidasikan menjadi:')
bullet('Functional Specification Document (FSD) per modul + aturan bisnis final.')
bullet('Konfigurasi sistem (CoA→kategori, produk, tarif, cron, hak akses) pada Blueprint.')
bullet('Rencana migrasi saldo awal, jadwal go-live, dan skenario UAT.')
para(' ')
para('Total pertanyaan: %d. Dokumen ini bersifat living document — dapat ditambah sesuai temuan '
     'pada sesi diskusi antara IT, User Operasional, dan Accounting.' % COUNTER['n'],
     italic=True, color=GREY, size=9.5)

doc.save('../../docs/Daftar_Pertanyaan_SME_Komprehensif_Semua_Modul.docx')
print('SAVED OK; total questions:', COUNTER['n'])
