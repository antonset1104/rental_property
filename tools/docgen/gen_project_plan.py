# -*- coding: utf-8 -*-
"""Project Plan & Timeline — 1 Juli 2026 s.d. Go-Live 1 Januari 2027."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x39, 0x64)
GREEN = RGBColor(0x00, 0x6A, 0x4E)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = Inches(0.6); sec.right_margin = Inches(0.6)
sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6)


def shade(c, h):
    tcPr = c._tc.get_or_add_tcPr(); s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), h)
    tcPr.append(s)


def ct(c, t, bold=False, white=False, size=9, color=None, center=False):
    c.text = ''; p = c.paragraphs[0]
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color: r.font.color.rgb = color


def h1(t):
    p = doc.add_heading(level=1); r = p.add_run(t); r.font.color.rgb = NAVY; r.font.size = Pt(15); return p


def h2(t):
    p = doc.add_heading(level=2); r = p.add_run(t); r.font.color.rgb = GREEN; r.font.size = Pt(12.5); return p


def para(t, bold=False, italic=False, size=10.5, color=None, after=6):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p


def bullet(t, pre=None):
    p = doc.add_paragraph(style='List Bullet')
    if pre:
        rr = p.add_run(pre); rr.bold = True
    p.add_run(t); return p


def table(headers, rows, widths, hdr_color='1F3964', sizes=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htx in enumerate(headers):
        shade(t.rows[0].cells[i], hdr_color); ct(t.rows[0].cells[i], htx, bold=True, white=True, size=9)
        t.rows[0].cells[i].width = widths[i]
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            ct(cells[i], val, size=(sizes[i] if sizes else 8.5))
            cells[i].width = widths[i]
    doc.add_paragraph()
    return t


# ==================== COVER ====================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('PROJECT PLAN & TIMELINE')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Implementasi Property Management System — Odoo 19 Community\n'
               '1 Juli 2026 — Go-Live 1 Januari 2027')
r.font.size = Pt(13); r.bold = True; r.font.color.rgb = GREEN
doc.add_paragraph()
meta = doc.add_table(rows=7, cols=2); meta.style = 'Table Grid'
for i, (k, v) in enumerate([
    ('Proyek', 'Implementasi PMS Odoo 19 (addon rental_management + 26 modul companion)'),
    ('Periode', '1 Juli 2026 – 1 Januari 2027 (6 bulan)'),
    ('Go-Live', '1 Januari 2027 — bertepatan dengan awal tahun fiskal (Jan–Des)'),
    ('Status saat plan disusun', 'Pengembangan 26 modul selesai (validasi statis); belum smoke test pada instance Odoo 19 live'),
    ('Metodologi', 'Fase bertahap (waterfall dengan iterasi perbaikan pada SIT/UAT)'),
    ('Disusun oleh', 'System Analyst / Functional & Odoo Developer'),
    ('Tanggal', '10 Juli 2026'),
]):
    shade(meta.rows[i].cells[0], 'E7ECF5'); ct(meta.rows[i].cells[0], k, bold=True, size=10)
    ct(meta.rows[i].cells[1], v, size=10)
    meta.rows[i].cells[0].width = Inches(2.3); meta.rows[i].cells[1].width = Inches(7.5)
doc.add_page_break()

# ==================== 1. RINGKASAN ====================
h1('1. Ringkasan Eksekutif')
para('Proyek ini mengimplementasikan Property Management System berbasis Odoo 19 Community '
     '(addon rental_management + 26 modul companion yang telah dikembangkan) hingga go-live '
     'produksi pada 1 Januari 2027. Tanggal go-live dipilih bertepatan dengan awal tahun fiskal '
     '(Januari–Desember) sehingga saldo awal, budget, dan pelaporan Owners Statement dimulai '
     'bersih dari periode pertama tahun buku.')
bullet('Seluruh kode modul telah selesai dibangun dan lolos validasi statis (sintaks Python/XML/CSV); '
       'belum pernah dijalankan pada instance Odoo 19 live — karena itu fase awal difokuskan pada '
       'smoke test menyeluruh di staging.', pre='Posisi saat ini: ')
bullet('Jawaban kuesioner SME (117 pertanyaan, area A–AF) menjadi prasyarat konfigurasi; beberapa '
       'keputusan kritikal (arah integrasi CORETAX) masih terbuka dan dijadwalkan final pada Fase 1.', pre='Dependensi: ')
bullet('Paralel run dua siklus billing penuh (November & Desember 2026) sebelum cut-over.', pre='Mitigasi utama: ')

h2('1.1 Sasaran Proyek')
bullet('Sistem produksi berjalan stabil per 1 Januari 2027 dengan seluruh modul terpasang & terkonfigurasi.')
bullet('Owners Statement Desember 2026 (hasil paralel run) tervalidasi cocok dengan sistem lama (MRI).')
bullet('Data master & saldo awal termigrasi lengkap dan tervalidasi (tenant balance, deposit, budget, aset).')
bullet('User operasional & accounting terlatih dan lulus UAT (79+ skenario, target 100% eksekusi, ≥95% pass).')

# ==================== 2. LINGKUP ====================
h1('2. Lingkup')
table(['Termasuk (In-Scope)', 'Tidak Termasuk (Out-of-Scope)'],
      [['Instalasi & konfigurasi Odoo 19 CE + rental_management + 26 modul companion pada server staging & produksi',
        'Kustomisasi baru di luar 26 modul yang sudah dibangun (masuk change request)'],
       ['Konfigurasi: CoA & Report Category, owners %, analytic, trust/deposit, produk layanan, tarif, cron, hak akses, portal',
        'Integrasi host-to-host bank / virtual account / payment gateway / POS (fase lanjutan)'],
       ['Migrasi data dari MRI: master properti/unit/tenant/kontrak + saldo awal (tenant balance, deposit, budget, aset & akumulasi)',
        'Pembersihan (cleansing) data sumber di sisi MRI — tanggung jawab tim data user'],
       ['SIT, UAT (79+ skenario), pelatihan user, paralel run 2 bulan, cut-over, dukungan hypercare',
        'Perubahan proses bisnis di luar yang terdokumentasi pada Blueprint'],
       ['Validasi XML CORETAX terhadap XSD resmi DJP & uji unggah ke portal CORETAX',
        'Pengurusan administrasi perpajakan itu sendiri (tetap oleh tim pajak)']],
      [Inches(4.9), Inches(4.9)], sizes=[8.5, 8.5])
doc.add_page_break()

# ==================== 3. FASE & TIMELINE ====================
h1('3. Fase Proyek & Timeline')
para('Enam fase dalam 6 bulan. Tanggal adalah target; pergeseran dikelola lewat manajemen risiko (Bagian 7).')
table(['Fase', 'Periode', 'Aktivitas Utama', 'Deliverable / Exit Criteria'],
      [['F1. Persiapan & Finalisasi Kebutuhan',
        '1 Jul – 31 Jul 2026',
        'Kick-off; pengisian & konsolidasi kuesioner SME (117 pertanyaan); finalisasi keputusan terbuka '
        '(arah integrasi CORETAX, mapping akun→Report Category dari Accounting); penyiapan server staging; '
        'freeze lingkup.',
        'Kuesioner SME terisi & disetujui; keputusan CORETAX final; staging siap; scope freeze ditandatangani.'],
       ['F2. Instalasi, Smoke Test & Stabilisasi',
        '1 Agu – 31 Agu 2026',
        'Instal Odoo 19 CE + rental_management + 26 modul di staging (urutan dependensi per Blueprint §4); '
        'smoke test seluruh titik sensitif versi (inherit view, analytic, portal); perbaikan bug instalasi; '
        'validasi XML CORETAX vs XSD DJP.',
        'Semua modul terpasang tanpa error; daftar temuan smoke test ditutup; XML CORETAX lolos XSD.'],
       ['F3. Konfigurasi & Migrasi Data (siklus 1)',
        '1 Sep – 30 Sep 2026',
        'Konfigurasi penuh per Blueprint §5 (12 subseksi); template import & trial migrasi data MRI '
        '(master + saldo awal); rekonsiliasi hasil trial; SIT alur ujung-ke-ujung.',
        'Konfigurasi selesai & didokumentasikan; trial migrasi rekonsil 100% pada sampel; SIT lulus.'],
       ['F4. UAT & Pelatihan',
        '1 Okt – 31 Okt 2026',
        'Pelatihan user operasional & accounting (berbasis Panduan Pengguna 13 BAB); eksekusi UAT 79+ '
        'skenario; perbaikan defect; regresi; sign-off UAT.',
        'UAT 100% dieksekusi, ≥95% pass, tanpa defect kritikal terbuka; berita acara sign-off UAT.'],
       ['F5. Paralel Run & Persiapan Cut-over',
        '1 Nov – 31 Des 2026',
        'Migrasi data final (cut-off 31 Okt); paralel run siklus billing November & Desember di dua sistem; '
        'bandingkan Owners Statement vs MRI; go/no-go review pertengahan Desember; migrasi delta & '
        'penguncian; gladi cut-over.',
        'Dua siklus paralel cocok (selisih = 0 atau terjelaskan); keputusan GO tertulis; runbook cut-over final.'],
       ['F6. Go-Live & Hypercare',
        '1 Jan 2027 (+ 4 minggu hypercare)',
        'Cut-over produksi 1 Januari 2027; MRI menjadi baca-saja; pendampingan intensif; tutup buku Januari '
        'sebagai siklus produksi pertama; serah terima ke operasi rutin.',
        'Sistem produksi stabil; billing Januari & Owners Statement Januari terbit dari sistem baru.']],
      [Inches(1.8), Inches(1.2), Inches(4.0), Inches(2.8)], sizes=[8.5, 8.5, 8, 8])
doc.add_page_break()

# ==================== 4. GANTT ====================
h1('4. Gantt Chart (Ringkas)')
MONTHS = ['Jul 26', 'Agu 26', 'Sep 26', 'Okt 26', 'Nov 26', 'Des 26', 'Jan 27']
GANTT = [
    ('F1 Persiapan & Kebutuhan',      [1, 0, 0, 0, 0, 0, 0]),
    ('F2 Instalasi & Smoke Test',     [0, 1, 0, 0, 0, 0, 0]),
    ('F3 Konfigurasi & Migrasi 1',    [0, 0, 1, 0, 0, 0, 0]),
    ('F4 UAT & Pelatihan',            [0, 0, 0, 1, 0, 0, 0]),
    ('F5 Paralel Run & Cut-over Prep', [0, 0, 0, 0, 1, 1, 0]),
    ('F6 Go-Live & Hypercare',        [0, 0, 0, 0, 0, 0, 1]),
]
g = doc.add_table(rows=1, cols=1 + len(MONTHS)); g.style = 'Table Grid'; g.alignment = WD_TABLE_ALIGNMENT.CENTER
ct(g.rows[0].cells[0], 'Fase', bold=True, white=True, size=9)
shade(g.rows[0].cells[0], '1F3964')
g.rows[0].cells[0].width = Inches(2.6)
for i, m in enumerate(MONTHS):
    c = g.rows[0].cells[1 + i]; shade(c, '1F3964'); ct(c, m, bold=True, white=True, size=8.5, center=True)
    c.width = Inches(1.0)
for name, cells_on in GANTT:
    row = g.add_row().cells
    ct(row[0], name, bold=True, size=8.5)
    row[0].width = Inches(2.6)
    for i, on in enumerate(cells_on):
        c = row[1 + i]; c.width = Inches(1.0)
        if on:
            shade(c, '006A4E'); ct(c, '█', white=True, size=8.5, center=True)
        else:
            ct(c, '', size=8.5)
doc.add_paragraph()
para('Milestone (♦): M1 Scope freeze — 31 Jul • M2 Staging stabil — 31 Agu • M3 SIT lulus — 30 Sep • '
     'M4 UAT sign-off — 31 Okt • M5 Keputusan GO — ±15 Des • M6 GO-LIVE — 1 Jan 2027.',
     bold=True, color=NAVY, size=10)
doc.add_page_break()

# ==================== 5. ORGANISASI ====================
h1('5. Organisasi Proyek & Tanggung Jawab')
table(['Peran', 'Personil (diisi)', 'Tanggung Jawab Utama'],
      [['Project Sponsor', '', 'Keputusan go/no-go, penyelesaian eskalasi, penandatangan scope freeze.'],
       ['Project Manager (IT)', '', 'Kendali jadwal/lingkup/risiko, laporan mingguan, koordinasi lintas tim.'],
       ['System Analyst / Functional', '', 'Konfigurasi sistem, dokumentasi, pendampingan UAT & pelatihan.'],
       ['Odoo Developer / Technical', '', 'Instalasi, perbaikan temuan smoke test/SIT/UAT, skrip migrasi.'],
       ['Key User Operasional', '', 'Jawaban SME, eksekusi UAT operasional, paralel run leasing/billing.'],
       ['Key User Accounting/Finance', '', 'Mapping CoA→Report Category, saldo awal, rekonsiliasi paralel run, UAT keuangan.'],
       ['Tim Pajak', '', 'Validasi CORETAX (kode transaksi, XSD, uji unggah), keputusan arah integrasi.'],
       ['Tim Data (MRI)', '', 'Ekstraksi & pembersihan data sumber, rekonsiliasi migrasi.'],
       ['IT Infrastruktur', '', 'Server staging & produksi, backup, SMTP, monitoring, keamanan.']],
      [Inches(2.2), Inches(1.8), Inches(5.8)], sizes=[9, 9, 8.5])

h2('5.1 Ritme Rapat')
bullet('Mingguan: status proyek (PM + perwakilan tiap tim, 30–45 menit).', pre='Weekly: ')
bullet('Per akhir fase: gate review terhadap exit criteria fase (keputusan lanjut/ulang).', pre='Gate: ')
bullet('Desember: go/no-go meeting khusus dengan sponsor (±15 Des 2026).', pre='Go/No-Go: ')
doc.add_page_break()

# ==================== 6. MIGRASI & CUT-OVER ====================
h1('6. Strategi Migrasi Data & Cut-over')
table(['Item Migrasi', 'Sumber', 'Metode', 'Validasi'],
      [['Master properti, unit, tenant, kontrak aktif', 'MRI / Excel', 'Template import Odoo (CSV/XLSX)', 'Hitung jumlah record & spot check per properti'],
       ['Saldo tenant berjalan (per kontrak)', 'MRI', 'Jurnal saldo awal per tenant', 'Total = laporan arrears MRI per cut-off'],
       ['Security deposit (titipan)', 'MRI', 'Jurnal liabilitas deposit per tenant', 'Total = daftar deposit MRI'],
       ['Budget 2026 sisa tahun + budget 2027', 'Excel PM', 'Import property.budget per akun/bulan', 'Total per properti = lembar budget disetujui'],
       ['Fixed asset & akumulasi penyusutan', 'MRI / daftar aset', 'Import register aset + saldo akumulasi', 'NBV total = neraca per cut-off'],
       ['Saldo GL awal (trust, bank, piutang, dsb.)', 'Neraca MRI', 'Jurnal opening balance', 'Trial balance seimbang & = neraca MRI']],
      [Inches(2.6), Inches(1.4), Inches(2.6), Inches(3.2)], sizes=[8.5, 8.5, 8.5, 8.5])
h2('6.1 Prinsip Cut-over')
bullet('Cut-off data final 31 Oktober 2026; transaksi Nov–Des dicatat ganda (paralel run) di kedua sistem.')
bullet('Selama paralel run, sistem lama (MRI) tetap menjadi sistem pencatatan resmi; Odoo menjadi pembanding.')
bullet('Keputusan GO (±15 Des) mensyaratkan: dua siklus billing cocok, UAT sign-off, tidak ada defect kritikal terbuka, runbook cut-over teruji.')
bullet('Per 1 Januari 2027 posisi dibalik: Odoo resmi, MRI baca-saja (dibekukan).')
bullet('Rollback plan: bila kegagalan kritikal pada minggu pertama Januari, pencatatan kembali ke MRI dan Odoo diperlakukan sebagai paralel — keputusan pada PM + sponsor.')

# ==================== 7. RISIKO ====================
doc.add_page_break()
h1('7. Risiko & Mitigasi')
table(['Risiko', 'Dampak', 'Kemungkinan', 'Mitigasi'],
      [['Kode belum pernah jalan di Odoo live — temuan smoke test lebih banyak dari perkiraan',
        'Tinggi', 'Sedang',
        'Seluruh Agustus didedikasikan untuk smoke test & stabilisasi; buffer perbaikan di September; daftar titik sensitif versi sudah terdokumentasi.'],
       ['Keputusan arah integrasi CORETAX terlambat',
        'Tinggi', 'Sedang',
        'Dijadwalkan final di F1 (Juli); bila mundur, modul lain tetap jalan — CORETAX di-uji terpisah sebelum paralel run.'],
       ['Kualitas data MRI buruk / mapping tidak lengkap',
        'Tinggi', 'Sedang',
        'Trial migrasi sejak September (2 siklus sebelum final); tanggung jawab cleansing di tim data user dengan tenggat jelas.'],
       ['Ketersediaan key user terbatas saat UAT/paralel run',
        'Sedang', 'Tinggi',
        'Jadwal UAT dikunci sejak kick-off; komitmen alokasi waktu ditandatangani sponsor; sesi pelatihan direkam.'],
       ['Mapping akun→Report Category dari Accounting terlambat (prasyarat Owners Statement)',
        'Tinggi', 'Sedang',
        'Sudah diminta lewat kuesioner SME; tenggat akhir Juli; template mapping disiapkan tim proyek untuk tinggal diisi.'],
       ['Beban ganda saat paralel run (input 2 sistem)',
        'Sedang', 'Tinggi',
        'Lingkup paralel dibatasi transaksi inti (billing, penerimaan, pengeluaran); dukungan input oleh tim proyek.'],
       ['Perubahan lingkup di tengah proyek',
        'Sedang', 'Sedang',
        'Scope freeze akhir Juli; permintaan baru masuk change request untuk fase pasca go-live.']],
      [Inches(3.0), Inches(0.9), Inches(1.1), Inches(4.8)], sizes=[8.5, 8.5, 8.5, 8])

# ==================== 8. KRITERIA GO-LIVE ====================
h1('8. Kriteria Go-Live (Go/No-Go, ±15 Desember 2026)')
table(['#', 'Kriteria', 'Target'],
      [['1', 'Eksekusi UAT', '100% skenario dieksekusi, ≥95% pass, 0 defect kritikal terbuka'],
       ['2', 'Paralel run billing Nov & Des', 'Nilai tagihan, penerimaan & arrears cocok dengan MRI (selisih 0 / terjelaskan & disetujui Finance)'],
       ['3', 'Owners Statement paralel', 'Statement Nov (dan pre-close Des) tervalidasi oleh Accounting & disetujui'],
       ['4', 'Migrasi data', 'Rekonsiliasi 100% pada master & saldo awal; berita acara migrasi ditandatangani'],
       ['5', 'CORETAX', 'XML lolos XSD DJP & uji unggah berhasil (atau prosedur manual disetujui bila arah integrasi berubah)'],
       ['6', 'Kesiapan user', 'Seluruh user inti terlatih; panduan & runbook dibagikan'],
       ['7', 'Kesiapan infrastruktur', 'Produksi terpasang, backup otomatis teruji restore, SMTP & cron (22.00–05.00 WIB) berjalan'],
       ['8', 'Dukungan', 'Jadwal hypercare 4 minggu & jalur eskalasi disepakati']],
      [Inches(0.5), Inches(3.4), Inches(5.9)], sizes=[9, 9, 8.5])

para(' ')
para('Catatan: Rencana ini disusun berdasarkan status kode per 10 Juli 2026 (selesai dikembangkan, '
     'lolos validasi statis, belum smoke test pada Odoo 19 live). Durasi F2 dapat memendek bila hasil '
     'smoke test baik — waktu sisa dialihkan menjadi buffer F3/F5. Dokumen ini hidup; versi terbaru '
     'selalu yang ada di repository.', italic=True, color=GREY, size=9.5)

doc.save('../../docs/Project_Plan_Timeline_Implementasi.docx')
print('SAVED OK')
