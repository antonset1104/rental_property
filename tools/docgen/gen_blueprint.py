# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x39, 0x64)
GREEN = RGBColor(0x00, 0x6A, 0x4E)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)


def shade(c, h):
    tcPr = c._tc.get_or_add_tcPr(); s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), h)
    tcPr.append(s)


def ct(c, t, bold=False, white=False, size=8.5, color=None):
    c.text = ''; p = c.paragraphs[0]; r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color: r.font.color.rgb = color


def h1(t):
    p = doc.add_heading(level=1); r = p.add_run(t); r.font.color.rgb = NAVY; r.font.size = Pt(15); return p


def h2(t):
    p = doc.add_heading(level=2); r = p.add_run(t); r.font.color.rgb = GREEN; r.font.size = Pt(12.5); return p


def para(t, bold=False, italic=False, size=10.5, color=None, after=6):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); return p


def bullet(t, pre=None):
    p = doc.add_paragraph(style='List Bullet')
    if pre:
        rr = p.add_run(pre); rr.bold = True
    p.add_run(t); return p


def steps(items):
    for s in items:
        doc.add_paragraph(style='List Number').add_run(s)


def table(headers, rows, widths, hdr_color='1F3964', sizes=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htx in enumerate(headers):
        shade(t.rows[0].cells[i], hdr_color); ct(t.rows[0].cells[i], htx, bold=True, white=True, size=9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            ct(cells[i], val, size=(sizes[i] if sizes else 8.5))
            cells[i].width = widths[i]
    doc.add_paragraph()
    return t


DIAG = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'docs', 'diagrams')


def figure(fname, caption, width=6.6):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture('%s/%s' % (DIAG, fname), width=Inches(width))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(10)


# ============ COVER ============
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('BLUEPRINT SISTEM\nProperty Management System – Odoo 19 Community')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Arsitektur • Katalog Modul • Manual Instalasi • Konfigurasi Sistem')
r.font.size = Pt(13); r.bold = True; r.font.color.rgb = GREEN
doc.add_paragraph()
meta = doc.add_table(rows=6, cols=2); meta.style = 'Table Grid'
for i, (k, v) in enumerate([
    ('Dokumen', 'System Blueprint, Installation & Configuration Manual'),
    ('Platform', 'Odoo 19 Community Edition'),
    ('Basis', 'addon rental_management v3.3.9 (TechKhedut)'),
    ('Lingkup', '26 modul companion + integrasi modul standar Odoo'),
    ('Disusun oleh', 'System Analyst / Functional & Odoo Developer'),
    ('Tanggal', '20 Juni 2026'),
]):
    shade(meta.rows[i].cells[0], 'E7ECF5'); ct(meta.rows[i].cells[0], k, bold=True, size=10)
    ct(meta.rows[i].cells[1], v, size=10)
    meta.rows[i].cells[0].width = Inches(1.8); meta.rows[i].cells[1].width = Inches(4.9)
doc.add_page_break()

# ============ 1. EXECUTIVE & ARCHITECTURE ============
h1('1. Ringkasan Eksekutif & Arsitektur')
para('Sistem ini adalah Property Management System (PMS) komprehensif yang dibangun di atas '
     'addon pihak ketiga rental_management (TechKhedut) untuk Odoo 19 Community. Fungsionalitas '
     'diperluas melalui 26 modul companion yang masing-masing bersifat opsional (decoupled) dan '
     'terintegrasi dengan modul standar Odoo: Accounting, Analytic, Purchase, Project, CRM, '
     'Portal, Maintenance, Product, Website, Mail.')
h2('1.1 Prinsip Arsitektur')
bullet('Companion, bukan fork: addon berlisensi tidak ditambal; semua ekstensi via modul terpisah.', pre='Non-invasif: ')
bullet('Tiap modul mengecek keberadaan field/model sebelum memakainya → dapat diinstal mandiri.', pre='Decoupled: ')
bullet('Semua transaksi keuangan ditautkan ke properti & otomatis ter-tag analytic saat posting.', pre='Terintegrasi: ')
bullet('Mengandalkan account.move / analytic / portal standar Odoo, bukan ledger paralel.', pre='Standar Odoo: ')
h2('1.2 Lapisan Sistem')
table(['Lapisan', 'Komponen'],
      [['Operasional Leasing', 'Kontrak, CRM, Casual Leasing, Handover, GTO, Meter, Parking, Vacancy, Lease Expiry, Guarantee, Insurance, Access/Visitor, E-sign, PPM'],
       ['Keuangan & Akuntansi', 'Owners Statement (9 laporan), Trust Accounting, Owner Remittance (multi-currency), Budget (+approval), Security Deposit, CAM, Rent Escalation, Dunning, Fixed Asset+Revaluation, Valuation, Procurement, CORETAX, Analytic'],
       ['Self-Service', 'Tenant Portal, Owner Portal'],
       ['Analitik', 'KPI Dashboard, Analytic Accounting, laporan standar Odoo'],
       ['Dokumen', 'Documents (Enterprise) / Documents CE (Community)']],
      [Inches(1.7), Inches(5.0)])
h2('1.3 Diagram Arsitektur')
figure('01-arsitektur-sistem.png',
       'Gambar 1. Arsitektur sistem berlapis: Lapisan Akses → 26 modul companion → '
       'addon rental_management → modul standar Odoo 19.')
figure('02-arsitektur-data.png',
       'Gambar 2. Arsitektur data — rantai atribusi properti hingga analytic & Owners Statement.')
figure('08-analytic-crosscutting.png',
       'Gambar 3. Integrasi cross-cutting: override account.move._post mencap analytic seluruh dokumen.')
doc.add_page_break()

# ============ 2. MODULE CATALOG ============
h1('2. Katalog Modul (26)')
MODULES = [
    ('rental_management_financial_report', 'Keuangan', 'Owners Statement 9 laporan, Trust, Remittance, Budget+approval, Deposit, Analytic stamping', 'Properties → Financial Reports', 'CoA→Report Category, Trust/Deposit accounts & journals, owners %'),
    ('rental_management_gto_meter', 'Leasing', 'GTO/Revenue sharing, Meter & utility recharge', 'Properties → Leasing Operations', 'GTO % di kontrak, produk GTO/utility, tarif meter'),
    ('rental_management_guarantee', 'Risiko', 'Bank/Insurance guarantee + alert expiry', 'Properties → Tenant Guarantees', 'reminder lead days; cron aktif'),
    ('rental_management_casual_leasing', 'Leasing', 'Sewa jangka pendek (booth/kios)', 'Properties → Casual Leases', 'produk Casual Leasing'),
    ('rental_management_handover', 'Operasional', 'Move-in/Fit-out/Move-out + checklist', 'Properties → Handovers', '—'),
    ('rental_management_portal', 'Self-Service', 'Tenant portal /my/contracts', '/my', 'grant portal access ke tenant'),
    ('rental_management_coretax', 'Pajak', 'e-Faktur & SPT CORETAX (11 ekspor XML)', 'Properties → CORETAX', 'TIN/IDTKU company & partner, kode produk'),
    ('rental_management_purchase', 'Procurement', 'PO↔properti/kontrak/maintenance → bill', 'Properties → Property Purchase Orders', 'aktifkan app Purchase'),
    ('rental_management_crm', 'Sales', 'Pipeline leasing, Create Lease Contract', 'CRM', 'team Property Leasing'),
    ('rental_management_project', 'Operasional', 'Fit-out → Project & Task', 'form Handover (Fit-out)', 'aktifkan app Project'),
    ('rental_management_documents', 'Dokumen', 'Folder Documents per properti', 'form Properti (Documents)', 'Enterprise; app Documents'),
    ('rental_management_asset', 'Aset', 'Fixed asset, depreciation, revaluation, sync L9', 'Properties → Fixed Assets', 'akun aset/akumulasi/beban/reserve & journal'),
    ('rental_management_document_ce', 'Dokumen', 'Lampiran properti via ir.attachment (Community)', 'form Properti (Files)', '—'),
    ('rental_management_owner_portal', 'Self-Service', 'Owner portal /my/properties + remittance', '/my', 'grant portal access ke owner'),
    ('rental_management_cam', 'Keuangan', 'CAM/Service charge: budget vs actual + apportion', 'Properties → CAM / Service Charge', 'produk Service Charge, total area, area tenant'),
    ('rental_management_rent_escalation', 'Leasing', 'Kenaikan sewa berkala otomatis (cron)', 'tab kontrak + Rent Escalation Log', 'tipe/nilai/frekuensi/next date di kontrak; cron'),
    ('rental_management_dashboard', 'Analitik', 'KPI: NOI, arrears, collection, WALE', 'Properties → KPI Dashboard', '—'),
    ('rental_management_dunning', 'Keuangan', 'Tangga dunning + late fee (cron)', 'Properties → Dunning Levels', 'level (hari), template email, produk late fee; cron'),
    ('rental_management_insurance', 'Risiko', 'Polis asuransi bangunan + alert expiry', 'Properties → Insurance Policies', 'reminder lead; cron'),
    ('rental_management_lease_expiry', 'Leasing', 'Daftar lease berakhir + reminder perpanjangan', 'Properties → Lease Expiry', 'reminder days; cron'),
    ('rental_management_vacancy', 'Leasing', 'Papan vacancy/availability per status', 'Properties → Vacancy / Availability', '—'),
    ('rental_management_ppm', 'Maintenance', 'PPM otomatis generate maintenance request (cron)', 'Properties → Preventive Maintenance', 'frekuensi/next date/SLA; cron'),
    ('rental_management_parking', 'Operasional', 'Bay parkir, alokasi & invoice', 'Properties → Parking', 'produk Parking, tarif'),
    ('rental_management_valuation', 'Aset', 'Valuasi pasar berkala', 'Properties → Property Valuations', '—'),
    ('rental_management_access', 'Operasional', 'Access/parking card + visitor log', 'Properties → Access & Visitors', '—'),
    ('rental_management_esign', 'Operasional', 'Tracking permintaan tanda tangan kontrak', 'Properties → Signature Requests', '—'),
]
table(['Modul', 'Area', 'Fitur Utama', 'Menu', 'Konfigurasi Kunci'],
      [[m[0], m[1], m[2], m[3], m[4]] for m in MODULES],
      [Inches(1.85), Inches(0.85), Inches(1.95), Inches(1.45), Inches(1.6)],
      sizes=[7.5, 7.5, 7.5, 7.5, 7.5])
doc.add_page_break()

# ============ 3. DATA ARCHITECTURE ============
h1('3. Arsitektur Data & Integrasi')
para('Rantai atribusi properti yang menyatukan seluruh transaksi keuangan:')
para('tenancy_id / sold_id / maintenance_request_id / property_manual_id  →  '
     'account.move.property_financial_id (stored)  →  Owners Statement (filter) + '
     'Analytic stamping pada account.move._post()', bold=True, color=NAVY)
bullet('Pendapatan: sewa, GTO, meter, casual lease, CAM, parking → out_invoice ber-tag properti.')
bullet('Beban: maintenance, vendor bill dari PO, penyusutan aset, late fee → ber-tag properti.')
bullet('Trust/Owner: owner remittance & security deposit posting account.move.')
bullet('Saat posting, baris income/expense otomatis dicap analytic account properti → laporan Analytic & Budget standar Odoo.')
para('Semua invoice yang dibuat modul-modul ini diberi tenancy_id atau property_manual_id, '
     'sehingga muncul otomatis di Owners Statement (Payment Details, Income & Expenditure) dan '
     'analytic per properti.', italic=True, color=GREY)
figure('03-siklus-hidup-tenant.png',
       'Gambar 4. Siklus hidup tenant end-to-end (lead → aktivasi → billing → move-out).')
doc.add_page_break()

# ============ 4. INSTALLATION MANUAL ============
h1('4. Manual Instalasi')
h2('4.1 Prasyarat')
bullet('Odoo 19 Community Edition (server berjalan, akses admin).', pre='Platform: ')
bullet('PostgreSQL 14+; Python 3.11+ dengan dependensi Odoo standar.', pre='Database: ')
bullet('Addon dasar rental_management (TechKhedut) v3.3.9 — wajib (lisensi OPL-1/berbayar).', pre='Base addon: ')
bullet('Modul standar diaktifkan: Accounting, Purchase, Project, CRM, Portal, Maintenance, Website.', pre='Apps Odoo: ')
bullet('Opsional Enterprise: app Documents (hanya untuk rental_management_documents; gunakan rental_management_document_ce di Community).', pre='Enterprise: ')
bullet('Worker cron Odoo aktif (untuk escalation, dunning, PPM, expiry alerts) & outgoing mail server (untuk dunning).', pre='Infra: ')

h2('4.2 Langkah Instalasi')
steps([
    'Salin folder addon dasar rental_management dan ke-26 folder rental_management_* ke addons_path Odoo (mis. /opt/odoo/addons atau path kustom).',
    'Pastikan addons_path di odoo.conf memuat lokasi tersebut; restart service Odoo.',
    'Aktifkan Developer Mode (Settings → Developer Tools) lalu Apps → Update Apps List.',
    'Install modul standar Odoo terlebih dahulu: Accounting, Purchase, Project, CRM, Portal, Maintenance.',
    'Install addon dasar: rental_management.',
    'Install fondasi keuangan: rental_management_financial_report (disarankan lebih dulu karena modul lain memanfaatkan field property_manual_id & analytic).',
    'Install modul lain sesuai kebutuhan (semua decoupled). Untuk paket lengkap, install seluruh 26 modul.',
    'Verifikasi tidak ada error pada log saat instalasi; periksa menu Properties bertambah.',
])
h2('4.3 Urutan Dependensi (ringkas)')
table(['Grup', 'Modul', 'Prasyarat'],
      [['Inti', 'rental_management', 'account, crm, maintenance, website, hr, mail'],
       ['Fondasi keuangan', 'rental_management_financial_report', 'rental_management, account'],
       ['Keuangan lanjutan', 'cam, dunning, asset, coretax, parking, casual_leasing, gto_meter, valuation', 'rental_management (+account)'],
       ['Operasional', 'handover → project, ppm, access, esign, guarantee, insurance, lease_expiry, vacancy, rent_escalation', 'rental_management (+mail/project)'],
       ['Self-service', 'portal (tenant), owner_portal', 'portal (+financial_report)'],
       ['Dokumen', 'documents (Enterprise) / document_ce (Community)', 'documents / rental_management'],
       ['Analitik', 'dashboard', 'rental_management, account']],
      [Inches(1.4), Inches(3.6), Inches(1.7)], sizes=[8, 8, 8])
doc.add_page_break()

# ============ 5. CONFIGURATION ============
h1('5. Konfigurasi Sistem (Detail)')
para('Bagian ini menjabarkan konfigurasi tingkat-field per area. Ikuti urutan 5.1 → 5.12. '
     'Tanda (Wajib) menunjukkan setting yang harus diisi agar laporan/otomasi berfungsi; (Opsional) '
     'dapat dilewati bila modul terkait tidak dipakai.', italic=True, color=GREY)

# ---- 5.1 Company & Akuntansi dasar ----
h2('5.1 Perusahaan & Akuntansi Dasar (Wajib)')
steps([
    'Settings → Companies: lengkapi data perusahaan (nama legal, NPWP, alamat, mata uang fungsional).',
    'Accounting → Configuration → Settings: aktifkan "Analytic Accounting" dan (bila perlu) "Multi-Currencies" serta "Multi-Companies".',
    'Pastikan Fiscal Localization / Chart of Accounts sudah terpasang sebelum memetakan Report Category.',
    'Accounting → Configuration → Journals: siapkan jurnal Sales, Purchase, Bank, Miscellaneous; tambahkan jurnal khusus bila perlu (Trust, Remittance, Deposit, Depreciation).',
])

# ---- 5.2 Report Category mapping ----
h2('5.2 Pemetaan Akun → Report Category Owners Statement (Wajib)')
para('Owners Statement mengelompokkan baris GL berdasarkan field "Property Report Category" '
     '(model property.financial.category) pada tiap akun. Tanpa pemetaan ini, angka tidak masuk '
     'section yang benar. Buat kategori di Properties → Financial → Report Categories lalu set '
     'pada akun (Accounting → Chart of Accounts → buka akun → field Property Report Category).')
table(['Section Laporan', 'Tipe', 'Contoh Akun GL', 'Catatan'],
      [['Rental Income', 'Income', '4-xxxx Pendapatan Sewa', 'sewa pokok kontrak'],
       ['Tenant Recharge Income', 'Income', '4-xxxx Recharge Utilitas/CAM', 'meter, CAM, service charge'],
       ['Other / GTO Income', 'Income', '4-xxxx Pendapatan GTO/Parkir', 'GTO, parkir, casual lease'],
       ['Property Expenses', 'Expense', '6-xxxx Beban O&M, utilitas', 'maintenance, utilitas, kebersihan'],
       ['Management Fee', 'Expense', '6-xxxx Management Fee', 'fee pengelola'],
       ['Capital Expenditure', 'Capital', '1-xxxx CWIP / Aset', 'pengeluaran kapital (di bawah Net Cash)'],
       ['GST/PPN Control', 'Tax', '2-xxxx PPN Keluaran/Masukan', 'untuk GST Reconciliation']],
      [Inches(1.7), Inches(0.8), Inches(2.4), Inches(1.8)], sizes=[8, 8, 8, 8])
bullet('Sub-group: gunakan field parent pada kategori untuk membentuk hierarki section/sub-group (struktur MR-style).')
bullet('Akun trust/deposit/remittance TIDAK dipetakan ke Income/Expense (lihat 5.4–5.5) agar tidak terhitung sebagai pendapatan.')

# ---- 5.3 Properti: owners, manager, analytic ----
h2('5.3 Konfigurasi per Properti (Wajib)')
steps([
    'Buka properti → tab "Owners & Financial".',
    'Owners: tambahkan baris pemilik + ownership % (total harus 100%). Isi Property Manager, telepon, faks (tampil di kop Owners Statement).',
    'Klik "Create Analytic Account" → membuat account.analytic.account khusus properti (dipakai stamping otomatis).',
    'Isi data fisik (luas/area) bila memakai CAM apportionment by area.',
])

# ---- 5.4 Trust Accounting ----
h2('5.4 Trust Accounting & Owner Remittance (Wajib bila pakai trust)')
table(['Field (per properti)', 'Tipe Akun', 'Fungsi'],
      [['Trust Bank Account', 'Bank/Asset', 'rekening penampung dana milik owner'],
       ['Owners Remittance Account', 'Liability/Payable', 'utang ke owner saat remittance'],
       ['Remittance Journal', 'Journal (Misc/Bank)', 'jurnal posting remittance']],
      [Inches(2.1), Inches(1.7), Inches(3.1)], sizes=[8.5, 8.5, 8.5])
bullet('Posting remittance: Dr Owners Remittance / Cr Trust Bank. Multi-currency: pilih Remittance Currency; konversi ke mata uang perusahaan pada kurs tanggal remittance (amount_currency tercatat).', pre='Posting: ')
bullet('Roll-forward: Opening Trust + Net Cash − Remittances = Closing Trust (tampil di Performance Summary Cash & Trust).', pre='Aritmetika: ')

# ---- 5.5 Security Deposit ----
h2('5.5 Security Deposit (Wajib bila pakai deposit)')
table(['Field (per properti)', 'Tipe Akun', 'Fungsi'],
      [['Deposit Liability Account', 'Liability', 'deposit sebagai kewajiban (bukan income)'],
       ['Forfeiture Income Account', 'Income', 'pengakuan deduction/forfeit'],
       ['Deposit Journal', 'Journal', 'jurnal held/refund/deduction']],
      [Inches(2.2), Inches(1.4), Inches(3.3)], sizes=[8.5, 8.5, 8.5])
bullet('Alur: Mark as Held (Dr Bank/Cr Deposit Liability) → Deduction (Dr Liability/Cr Forfeiture Income) → Refund (Dr Liability/Cr Bank). Saldo muncul di kolom Sec Dep Bal pada Tenant Balances.')

# ---- 5.6 Produk layanan ----
h2('5.6 Produk Layanan (petakan income account-nya)')
table(['Produk', 'Modul', 'Income Account → Report Category'],
      [['Percentage / GTO Rent', 'gto_meter', 'Other / GTO Income'],
       ['Utility Recharge', 'gto_meter', 'Tenant Recharge Income'],
       ['Casual Leasing', 'casual_leasing', 'Other / GTO Income'],
       ['Service Charge / CAM', 'cam', 'Tenant Recharge Income'],
       ['Parking Rental', 'parking', 'Other / GTO Income'],
       ['Late Payment Fee', 'dunning', 'Other Income']],
      [Inches(2.0), Inches(1.4), Inches(3.3)], sizes=[8.5, 8.5, 8.5])
para('Buka tiap produk (otomatis dibuat saat install) → tab Accounting → set Income Account, lalu '
     'pastikan akun tersebut sudah dipetakan ke Property Report Category (5.2).', italic=True, color=GREY)

# ---- 5.7 GTO & Meter ----
h2('5.7 GTO / Revenue Sharing & Meter')
bullet('Di kontrak (tenancy) tab GTO: aktifkan GTO, pilih tipe (Higher-of / Base+Overage / Pure %), isi Turnover % dan Breakpoint (untuk Base+Overage), pilih produk GTO.', pre='GTO: ')
bullet('Meter: daftarkan per properti (jenis listrik/air/gas), pilih produk recharge & tarif per unit; reading berikut mengambil previous otomatis.', pre='Meter: ')

# ---- 5.8 CAM ----
h2('5.8 CAM / Service Charge')
steps([
    'Set produk Service Charge & income account-nya.',
    'Isi Total Area properti dan area tiap tenant (untuk apportionment by area).',
    'Buat pool biaya CAM per periode (budget), masukkan actual; selisih untuk rekonsiliasi.',
    'Apportion → tinjau alokasi per tenant → Create Invoices.',
])

# ---- 5.9 Fixed Asset ----
h2('5.9 Fixed Asset & Depreciation')
table(['Field', 'Fungsi'],
      [['Asset / Gross Value Account', 'nilai perolehan aset'],
       ['Accumulated Depreciation Account', 'akumulasi penyusutan (Cr)'],
       ['Depreciation Expense Account', 'beban penyusutan (Dr)'],
       ['Revaluation Reserve Account', 'cadangan revaluasi (naik/turun)'],
       ['Depreciation Journal', 'jurnal posting penyusutan & revaluasi'],
       ['Method / Number / Period', 'Garis Lurus atau Saldo Menurun; jumlah & panjang periode']],
      [Inches(2.6), Inches(4.1)], sizes=[8.5, 8.5])
bullet('Compute Depreciation → Confirm → Post Due (atau cron). Revaluation: tambah baris +/- → Post Revaluations (board sisa dihitung ulang). Sync CORETAX L9 bila modul coretax aktif.')

# ---- 5.10 CORETAX ----
h2('5.10 CORETAX (Lokalisasi Pajak Indonesia)')
steps([
    'Company & tiap pelanggan (Contacts → tab CORETAX): isi TIN/NPWP (16 digit), ID TKU, NIK bila perlu.',
    'Produk: set CORETAX Type (A=barang / B=jasa) + Unit Code; sewa = Service/B.',
    'Invoice pelanggan → tab CORETAX e-Faktur: cek Transaction Code (default 04), Additional Info, Facility Stamp bila ada fasilitas.',
    'Vendor credit note (retur): isi Original Faktur No. & tanggal untuk Retur PM.',
    'Register SPT: Properties → CORETAX → SPT Registers untuk L9, L3B, L11A (Uncollectible/NonPerforming/Promosi/Entertainment), L10A — isi per Tax Year.',
    'Ekspor: CORETAX → e-Faktur / SPT Export → pilih jenis + periode/Tax Year → Export XML.',
    'WAJIB validasi XML hasil terhadap XSD resmi DJP sebelum diunggah ke portal CORETAX.',
])

# ---- 5.11 Cron ----
h2('5.11 Scheduled Actions (Cron) — aktifkan & atur interval')
para('Settings → Technical → Scheduled Actions. Aktifkan dan sesuaikan interval (umumnya harian, '
     'di luar jam sibuk). Pastikan worker cron Odoo berjalan.', italic=True, color=GREY)
table(['Scheduled Action', 'Modul', 'Interval saran', 'Fungsi'],
      [['Apply Due Rent Escalations', 'rent_escalation', 'Harian', 'naikkan sewa pada tanggal jatuh tempo'],
       ['Run Dunning', 'dunning', 'Harian', 'kirim reminder + late fee bertingkat'],
       ['Generate Preventive Maintenance', 'ppm', 'Harian', 'buat maintenance request terjadwal'],
       ['Guarantee Check Expiry', 'guarantee', 'Harian', 'auto-expire + reminder activity'],
       ['Insurance Check Expiry', 'insurance', 'Harian', 'auto-expire + reminder activity'],
       ['Lease Expiry / Renewal Reminder', 'lease_expiry', 'Harian', 'aktivitas perpanjangan 30/90 hari'],
       ['Post Due Depreciation', 'asset', 'Harian/Bulanan', 'posting baris penyusutan jatuh tempo']],
      [Inches(2.3), Inches(1.2), Inches(1.1), Inches(2.1)], sizes=[8, 8, 8, 8])
bullet('Dunning: definisikan level di Properties → Dunning Levels (hari overdue, template email, late fee on/off) sebelum mengaktifkan cron.', pre='Dunning: ')
bullet('Escalation: set tipe (% / amount), nilai, frekuensi, dan Next Escalation Date di kontrak.', pre='Escalation: ')
bullet('PPM: set frekuensi, Next Date, dan SLA pada plan; cron membuat request & menggeser Next Date.', pre='PPM: ')

# ---- 5.12 Sequences, Portal, Email ----
h2('5.12 Penomoran, Portal, Email & Multi-Entitas')
bullet('Penomoran (ir.sequence) terbuat otomatis (Owner Remittance, Security Deposit, dll). Sesuaikan prefix/format di Settings → Technical → Sequences bila perlu.', pre='Sequence: ')
bullet('Tenant: Contacts → kontak → Action → Grant portal access. Owner: idem untuk pemilik (muncul di /my/properties).', pre='Portal: ')
bullet('Settings → Technical → Outgoing Mail Servers (SMTP) wajib aktif untuk email dunning/reminder; uji kirim test.', pre='Email: ')
bullet('Aktifkan Multi-Company & Multi-Currency di Settings bila portofolio lintas entitas/mata uang; set kurs di Accounting → Currencies.', pre='Multi: ')
bullet('Atur Hak Akses (Settings → Users): Accounting Manager untuk approval Budget & ekspor CORETAX; Portal untuk tenant/owner.', pre='Akses: ')
doc.add_page_break()

# ============ 6. SECURITY ============
h1('6. Keamanan & Hak Akses')
bullet('Mayoritas model: akses base.group_user (read/write/create/unlink).')
bullet('Approval budget & CORETAX manager: account.group_account_manager.')
bullet('Portal: base.group_portal dengan record rule (tenant hanya kontraknya; owner hanya propertinya).')
bullet('Nama properti diekspos ke portal via stored related agar model properti tidak terbuka penuh.')

# ============ 7. WORKFLOWS ============
h1('7. Alur Operasional Inti (End-to-End)')
steps([
    'Akuisisi tenant: Lead (CRM) → Create Lease Contract → aktivasi kontrak → Handover Move-in → catat Guarantee.',
    'Billing berkala: invoice sewa + GTO turnover + meter recharge + CAM + parking → semua ber-tag properti & analytic.',
    'Pengeluaran: Maintenance/PPM → Create PO → terima → vendor bill; penyusutan aset via cron.',
    'Penagihan: cron Dunning mengirim reminder + late fee untuk invoice menunggak.',
    'Tutup periode: cetak Owners Statement; buat & setujui Budget; Owner Remittance (multi-currency); ekspor CORETAX.',
    'Akhir sewa: Lease Expiry reminder → perpanjang atau Handover Move-out + make-good + selesaikan Security Deposit.',
])
figure('04-billing-penagihan.png',
       'Gambar 5. Alur billing bulanan & penagihan (generasi tagihan → posting → dunning → trust).')
figure('05-pengeluaran.png',
       'Gambar 6. Alur pengeluaran: procurement & maintenance → vendor bill ber-atribusi properti.')
figure('06-tutup-periode.png',
       'Gambar 7. Alur tutup periode: Owners Statement, rekonsiliasi trust & owner remittance.')
figure('07-fixed-asset.png',
       'Gambar 8. Alur fixed asset: depresiasi terjadwal & revaluation terhadap reserve.')

# ============ 8. MAINTENANCE ============
h1('8. Pemeliharaan, Backup & Go-Live')
bullet('Backup database PostgreSQL & filestore harian (otomatis).', pre='Backup: ')
bullet('Jalankan UAT (lihat UAT Tracker) di staging sebelum produksi; validasi titik API sensitif versi.', pre='UAT: ')
bullet('Pantau log cron (escalation/dunning/PPM/depreciation) setelah go-live.', pre='Monitoring: ')
bullet('Migrasi saldo awal: opening tenant balance, deposit, trust balance, budget, aset & penyusutan berjalan.', pre='Migrasi: ')
para(' ')
para('Catatan: seluruh modul lolos validasi sintaks Python/XML/CSV; WAJIB smoke test pada instance '
     'Odoo 19 Community (staging) sebelum produksi.', italic=True, color=GREY, size=9.5)

doc.save('../../docs/Blueprint_Installation_Configuration.docx')
print('BLUEPRINT SAVED')
